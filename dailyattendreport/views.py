from django.utils import timezone
from datetime import datetime
from django.conf import settings
from django.http import HttpResponse
from django.http import HttpResponseRedirect
from django.views.static import serve
from django.shortcuts import render
from django.contrib.auth.decorators import login_required
from django.contrib.auth.mixins import PermissionRequiredMixin
from django.contrib.auth.decorators import permission_required
from django.utils import translation
from django.utils.translation import ugettext as _
import django.db as db
from django.db import connection
from page.rules import *
import datetime
from django.utils import timezone
from system.models import ComZone
from django.http import JsonResponse
from hrms.settings import MEDIA_ROOT
from django.http import FileResponse
from docxtpl import DocxTemplate
from docx2pdf import convert
import xlwt
from os import path
from docx import Document
from docx.shared import Cm, Mm, Pt, Inches
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE


@permission_required('dailyattendreport.can_access_psn_slip_d1_report', login_url='/accounts/login/')
def GeneratePSNSlipD1(request, *args, **kwargs):    
    base_url = MEDIA_ROOT + '/monitoring/template/'
    template_name = base_url + 'PSN_SLIP.docx'
    file_name = request.user.username + "_PSN_SLIP"

    emp_type = 'D1'
    emp_id = kwargs['emp_id']
    period_option = kwargs['period']

    pay_sum_obj = None
    record = {}
    pay_sum_list = []
    employee_paysum_income_list = []
    employee_paysum_deduct_list = []

    error_message = ""

    print(emp_id, period_option)

    sql = "select a.*,b.dept_en,c.sts_en from employee as a left join com_department as b on a.emp_dept=b.dept_id "
    sql += "left join t_empsts as c on a.emp_status=c.sts_id where a.emp_id=" + str(emp_id) + " and a.emp_type='D1';"    
    # print("SQL 1: ", sql)

    period_last_3_digits = ""
    year_period_in_thai = ""
    paid_period = ""

    employee_info = None
    employee_paysum_list = []
    eps_paid_stat_text = '?'
    emp_id = ""
    emp_title_text = ""
    emp_full_name = ""
    emp_rank = ""
    emp_status = ""
    emp_dept = ""
    dept_en = ""
    dept_en_short = ""
    emp_join_date = ""
    emp_term_date = ""
    emp_acc_bank = ""
    emp_acc_no = ""
    pay_tax = ""

    eps_paid_stat_text = '?'
    # Gross Income
    eps_prd_in = ""
    # Net Income
    eps_prd_net = ""
    # YTD Income
    eps_ysm_in = ""
    # YTD Prov.Func
    eps_ysm_prv = ""
    # Total Deduct
    eps_prd_de = ""
    # Tax
    eps_prd_tax = ""
    # YTD Tax
    eps_ysm_tax = ""
    # YTD Social Security
    eps_ysm_soc = ""    

    try:                
        cursor = connection.cursor()
        cursor.execute(sql)
        employee_info = cursor.fetchone()
    except db.OperationalError as e:
        is_error = True
        error_message = "<b>Error: please send this error to IT team</b><br>" + str(e)
    except db.Error as e:
        is_error = True
        error_message = "<b>Error: please send this error to IT team</b><br>" + str(e)
    finally:
        cursor.close()

    if employee_info is not None:
        is_error = False
        error_message = ""

        emp_id = employee_info[0]

        emp_title = employee_info[1]
        if (emp_title==3):
            emp_title_text = 'นาย'
        elif (emp_title==4):
            emp_title_text = 'นางสาว'
        elif (emp_title==5):
            emp_title_text = 'นาง'
        elif (emp_title==129):
            emp_title_text = 'คุณ'
        else:
            emp_title_text = 'คุณ'

        emp_full_name = str(employee_info[4]).strip() + "  " + str(employee_info[5]).strip()
        emp_rank = employee_info[31]
        emp_status = employee_info[69]
        emp_dept = employee_info[28]
        dept_en = employee_info[68]
        if (dept_en.find('SP-Zone') != -1):
            dept_en_short = dept_en.replace("SP-Zone", "").lstrip()[0]
        elif (dept_en.find('ZONE H BPI') != -1):
            dept_en_short = "H"
        elif (dept_en.find('ZONE Control') != -1):
            dept_en_short = "CR"
        elif (dept_en.find('SP-Samui') != -1):
            dept_en_short = "SM"
        elif (dept_en.find('ZONE PHUKET') != -1):
            dept_en_short = "P"
        elif (dept_en.find('Zone Khon kean') != -1):
            dept_en_short = "K"
        elif (dept_en.find('BEM') != -1):
            dept_en_short = "BEM"
        elif (dept_en.find('ZONE I') != -1):
            dept_en_short = "I"
        elif (dept_en.find('ZONE R') != -1):
            dept_en_short = "R"
        elif (dept_en.find('Zone Songkhla') != -1):
            dept_en_short = "SK"
        else:
            dept_en_short = dept_en

        emp_join_date = "" if employee_info[39] is None else employee_info[39].strftime("%d/%m/%Y")
        emp_term_date = "" if employee_info[40] is None else employee_info[40].strftime("%d/%m/%Y")
        emp_acc_bank = employee_info[45]
        emp_acc_no = employee_info[43]        

        # TODO
        sql = "select distinct eps_prd_id from pay_sum where eps_prd_id='" + str(period_option) + "';"
        try:
            cursor = connection.cursor()
            cursor.execute(sql)
            period_obj = cursor.fetchall()
        except db.OperationalError as e:
            is_error = True
            error_message = "<b>Error: please send this error to IT team</b><br>" + str(e)
        except db.Error as e:
            is_error = True
            error_message = "<b>Error: please send this error to IT team</b><br>" + str(e)
        finally:
            cursor.close()
        table = "PAY_SUM"
        if period_obj is not None:
            if len(period_obj) > 0:
                # แสดงว่าเป็นงวดปัจจุบัน
                # table = "PAY_SUM"
                table = "R_PAYSLIP"
            else:
                # แสดงว่าเป็นงวดย้อนหลัง
                # table = "HIS_PAY_SUM"
                table = "R_HPAYSLIP"
        
        # PAYSUM / HIS_PAY_SUM
        '''
        sql = "SELECT a.*,b.pay_th,b.pay_tax FROM " + str(table) + " as A left join t_paytype as B on a.eps_pay_type=b.pay_type "
        sql += "where eps_prd_id='" + str(period_option) + "' and eps_emp_id=" + str(emp_id) + " "
        sql += "ORDER BY eps_pay_type, eps_pay_seq;"
        print("SQL 2: ", sql)
        '''

        # sql = "select prd_date_paid, emp_fname_th, emp_lname_th, dept_en, title_th, emp_acc_bank, emp_acc_no, emp_dept, emp_rank, emp_term_date "
        # sql += "eps_emp_id, eps_prd_id, eps_pay_type, eps_ysm_in, eps_ysm_prv, eps_ysm_soc, eps_ysm_tax, eps_prd_in, eps_prd_de, eps_prd_net, eps_prd_tax "
        # sql += ", pay_th, pay_inde, eps_comp, eps_percent, eps_wrk_day, eps_wrk_hr, pay_tax, eps_paid_stat "

        sql = "select * "
        sql += "FROM " + str(table) + " "
        sql += "where eps_prd_id='" + str(period_option) + "' and eps_emp_id=" + str(emp_id) + " "
        sql += "ORDER BY eps_prd_de, eps_prd_net, eps_prd_tax;"
        print("SQL 212: ", sql)

        
        '''
        # R_PAYSLIP / R_HPAYSLIP
        sql = "select eps_emp_id,eps_prd_id,eps_pay_type paytype,eps_prd_order,eps_inde,eps_emp_type,eps_emp_dept,eps_amt,eps_paid,eps_comp,"
        sql += "eps_percent,eps_pay_seq,eps_wrk_day,eps_wrk_hr,eps_ysm_in,eps_ysm_de,eps_ysm_bas,eps_ysm_otm,"
        sql += "eps_ysm_bon,eps_ysm_prv,eps_ysm_soc,eps_ysm_tax,eps_prd_stax,eps_prd_sntax,eps_prd_1soc,eps_prd_2soc,eps_prd_psoc,"
        sql += "eps_prd_in,eps_prd_de,eps_prd_net,eps_prd_tax,eps_paid_stat,eps_doc_no,eps_doc_date,upd_date,upd_by,upd_flag,pay_th,"
        sql += "pay_tax "
        sql += "from R_HPAYSLIP where eps_prd_id='D121031' and eps_emp_id=916 ORDER BY eps_pay_type;"
        '''

        employee_paysum_obj = None        
        record = {}
        try:
            cursor = connection.cursor()
            cursor.execute(sql)
            employee_paysum_obj = cursor.fetchall()
        except db.OperationalError as e:
            is_error = True
            error_message = "<b>Error: please send this error to IT team</b><br>" + str(e)
        except db.Error as e:
            is_error = True
            error_message = "<b>Error: please send this error to IT team</b><br>" + str(e)
        finally:
            cursor.close()

        if employee_paysum_obj is not None:
            if len(employee_paysum_obj) > 0:
                
                if table == "R_PAYSLIP":

                    row_count = 1
                    eps_paid_stat = ""
                    eps_paid_stat_text = ""                    

                    for item in employee_paysum_obj:

                        # print("AAA : ", item[2])

                        if item[2] == '001':
                            eps_paid_stat = item[67]
                            
                            if eps_paid_stat=='P':
                                eps_paid_stat_text = 'P : PAID'
                            elif eps_paid_stat=='H':
                                eps_paid_stat_text = 'H : HOLDING'
                            elif eps_paid_stat=='C':
                                eps_paid_stat_text = 'C : CHEQUE'
                            else:
                                eps_paid_stat_text = eps_paid_stat

                            # eps_prd_in = '{:,}'.format(item[27])
                            eps_prd_in = 0 if item[27] is None else '{:,}'.format(item[27])
                            # print("DEBUG : ", eps_prd_in)

                            # Net Income
                            # eps_prd_net = '{:,}'.format(item[29])
                            eps_prd_net = 0 if item[29] is None else '{:,}'.format(item[29])

                            # YTD Income
                            # eps_ysm_in = '{:,}'.format(item[14])
                            eps_ysm_in = 0 if item[14] is None else '{:,}'.format(item[14])

                            # YTD Prov.Func
                            # eps_ysm_prv = '{:,}'.format(item[19])                        
                            eps_ysm_prv = 0 if item[19] is None else '{:,}'.format(item[19])

                            # Total Deduct
                            # eps_prd_de = '{:,}'.format(item[28])
                            eps_prd_de = 0 if item[28] is None else '{:,}'.format(item[28])

                            # Tax
                            # eps_prd_tax = '{:,}'.format(item[30])
                            eps_prd_tax = 0 if item[30] is None else '{:,}'.format(item[30])

                            # YTD Tax
                            # eps_ysm_tax = '{:,}'.format(item[21])
                            eps_ysm_tax = 0 if item[21] is None else '{:,}'.format(item[21])

                            # YTD Social Security
                            # eps_ysm_soc = '{:,}'.format(item[20])
                            eps_ysm_soc = 0 if item[20] is None else '{:,}'.format(item[20])

                        '''
                        if (row_count == 1):
                            row_count = row_count + 1                        
                            
                            eps_paid_stat = item[31]
                            
                            if eps_paid_stat=='P':
                                eps_paid_stat_text = 'P : PAID'
                            elif eps_paid_stat=='H':
                                eps_paid_stat_text = 'H : HOLDING'
                            elif eps_paid_stat=='C':
                                eps_paid_stat_text = 'C : CHEQUE'
                            else:
                                eps_paid_stat_text = eps_paid_stat

                            # Gross Income
                            eps_prd_in = 0 if item[27] is None else '{:,}'.format(item[27])

                            # Net Income
                            eps_prd_net = 0 if item[29] is None else '{:,}'.format(item[29])

                            # YTD Income
                            eps_ysm_in = 0 if item[14] is None else '{:,}'.format(item[14])

                            # YTD Prov.Func
                            eps_ysm_prv = 0 if item[19] is None else '{:,}'.format(item[19])

                            # Total Deduct
                            eps_prd_de = 0 if item[28] is None else '{:,}'.format(item[28])

                            # Tax
                            eps_prd_tax = 0 if item[30] is None else '{:,}'.format(item[30])

                            # YTD Tax
                            eps_ysm_tax = 0 if item[21] is None else '{:,}'.format(item[21])

                            # YTD Social Security
                            eps_ysm_soc = 0 if item[20] is None else '{:,}'.format(item[20])
                        '''

                        eps_emp_id = item[0]
                        
                        eps_pay_type = item[2]
                        
                        pay_th = item[37]
                        
                        payment_type = str(eps_pay_type) + " " + str(pay_th)

                        income_or_deduct = item[7]

                        eps_inde = item[4]
                        
                        if (item[9] is not None):
                            eps_comp = item[9]
                        else:
                            eps_comp = ""

                        if (item[10] is not None):
                            eps_percent = item[10]
                        else:
                            eps_percent = ""

                        if (item[12] is not None):
                            eps_wrk_day = item[12]
                        else:
                            eps_wrk_day = ""

                        if (item[13] is not None):
                            eps_wrk_hr = item[13]
                        else:
                            eps_wrk_hr = ""

                        # PAY TAX
                        # if (item[38] == 1):
                        if (item[44] == 1):
                            pay_tax = 1
                        else:
                            pay_tax = 0

                        record = {
                            "eps_emp_id": eps_emp_id,
                            "payment_type": payment_type,
                            "eps_inde": eps_inde,
                            "income_or_deduct": income_or_deduct,
                            "eps_comp": eps_comp,
                            "eps_percent": eps_percent,
                            "eps_wrk_day": eps_wrk_day,
                            "eps_wrk_hr": eps_wrk_hr,
                            "eps_paid_stat": eps_paid_stat,
                            "pay_tax": pay_tax,
                        }

                        if (eps_inde!='S'):
                            employee_paysum_list.append(record)
                            if eps_inde=='I':
                                employee_paysum_income_list.append(record)

                            if eps_inde=='D':
                                employee_paysum_deduct_list.append(record)


                elif table == "R_HPAYSLIP":
                    row_count = 1
                    eps_paid_stat = ""
                    eps_paid_stat_text = ""

                    for item in employee_paysum_obj:

                        if item[38] == '001':
                            eps_paid_stat = item[67]
                            
                            if eps_paid_stat=='P':
                                eps_paid_stat_text = 'P : PAID'
                            elif eps_paid_stat=='H':
                                eps_paid_stat_text = 'H : HOLDING'
                            elif eps_paid_stat=='C':
                                eps_paid_stat_text = 'C : CHEQUE'
                            else:
                                eps_paid_stat_text = eps_paid_stat

                            # eps_prd_in = '{:,}'.format(item[27])
                            eps_prd_in = 0 if item[63] is None else '{:,}'.format(item[63])

                            # Net Income
                            # eps_prd_net = '{:,}'.format(item[29])
                            eps_prd_net = 0 if item[65] is None else '{:,}'.format(item[65])

                            # YTD Income
                            # eps_ysm_in = '{:,}'.format(item[14])
                            eps_ysm_in = 0 if item[50] is None else '{:,}'.format(item[50])

                            # YTD Prov.Func
                            # eps_ysm_prv = '{:,}'.format(item[19])                        
                            eps_ysm_prv = 0 if item[55] is None else '{:,}'.format(item[55])

                            # Total Deduct
                            # eps_prd_de = '{:,}'.format(item[28])
                            eps_prd_de = 0 if item[64] is None else '{:,}'.format(item[64])

                            # Tax
                            # eps_prd_tax = '{:,}'.format(item[30])
                            eps_prd_tax = 0 if item[66] is None else '{:,}'.format(item[66])

                            # YTD Tax
                            # eps_ysm_tax = '{:,}'.format(item[21])
                            eps_ysm_tax = 0 if item[57] is None else '{:,}'.format(item[57])

                            # YTD Social Security
                            # eps_ysm_soc = '{:,}'.format(item[20])
                            eps_ysm_soc = 0 if item[56] is None else '{:,}'.format(item[56])


                        '''
                        if (row_count == 1):
                            row_count = row_count + 1
                            
                            # eps_paid_stat = item[31]
                            eps_paid_stat = item[67]
                            
                            if eps_paid_stat=='P':
                                eps_paid_stat_text = 'P : PAID'
                            elif eps_paid_stat=='H':
                                eps_paid_stat_text = 'H : HOLDING'
                            elif eps_paid_stat=='C':
                                eps_paid_stat_text = 'C : CHEQUE'
                            else:
                                eps_paid_stat_text = eps_paid_stat

                            # Gross Income
                            print("GROSS INCOME : ", item[63])
                            # eps_prd_in = '{:,}'.format(item[27])
                            eps_prd_in = 0 if item[63] is None else '{:,}'.format(item[63])


                            # Net Income
                            # eps_prd_net = '{:,}'.format(item[29])
                            eps_prd_net = 0 if item[65] is None else '{:,}'.format(item[65])

                            # YTD Income
                            # eps_ysm_in = '{:,}'.format(item[14])
                            eps_ysm_in = 0 if item[50] is None else '{:,}'.format(item[50])

                            # YTD Prov.Func
                            # eps_ysm_prv = '{:,}'.format(item[19])                        
                            eps_ysm_prv = 0 if item[55] is None else '{:,}'.format(item[55])

                            # Total Deduct
                            # eps_prd_de = '{:,}'.format(item[28])
                            eps_prd_de = 0 if item[64] is None else '{:,}'.format(item[64])

                            # Tax
                            # eps_prd_tax = '{:,}'.format(item[30])
                            eps_prd_tax = 0 if item[58] is None else '{:,}'.format(item[58])

                            # YTD Tax
                            # eps_ysm_tax = '{:,}'.format(item[21])
                            eps_ysm_tax = 0 if item[57] is None else '{:,}'.format(item[57])

                            # YTD Social Security
                            # eps_ysm_soc = '{:,}'.format(item[20])
                            eps_ysm_soc = 0 if item[56] is None else '{:,}'.format(item[56])
                        '''

                        # eps_emp_id = item[0]
                        eps_emp_id = item[36]
                        
                        # eps_pay_type = item[2]
                        eps_pay_type = item[38]
                        
                        # pay_th = item[37]
                        pay_th = item[0]
                        
                        payment_type = str(eps_pay_type) + " " + str(pay_th)

                        # income_or_deduct = '{:,}'.format(item[26])
                        income_or_deduct = item[43]

                        # eps_inde = item[4]
                        eps_inde = item[4]
                        
                        '''
                        if (item[9] is not None):
                            eps_comp = item[9]
                        else:
                            eps_comp = ""
                        '''
                        if (item[45] is not None):
                            eps_comp = item[45]
                        else:
                            eps_comp = ""

                        '''
                        if (item[10] is not None):
                            eps_percent = item[10]
                        else:
                            eps_percent = ""
                        '''
                        if (item[46] is not None):
                            eps_percent = item[46]
                        else:
                            eps_percent = ""

                        '''                    
                        if (item[12] is not None):
                            eps_wrk_day = item[12]
                        else:
                            eps_wrk_day = ""
                        '''
                        if (item[48] is not None):
                            eps_wrk_day = item[48]
                        else:
                            eps_wrk_day = ""

                        
                        '''
                        if (item[13] is not None):
                            eps_wrk_hr = item[13]
                        else:
                            eps_wrk_hr = ""
                        '''                    
                        if (item[49] is not None):
                            eps_wrk_hr = item[49]
                        else:
                            eps_wrk_hr = ""


                        # PAY TAX
                        '''
                        if (item[38] == 1):
                            pay_tax = 1
                        else:
                            pay_tax = 0
                        '''
                        if (item[5] == 1):
                            pay_tax = 1
                        else:
                            pay_tax = 0


                        record = {
                            "eps_emp_id": eps_emp_id,
                            "payment_type": payment_type,
                            "eps_inde": eps_inde,
                            "income_or_deduct": income_or_deduct,
                            "eps_comp": eps_comp,
                            "eps_percent": eps_percent,
                            "eps_wrk_day": eps_wrk_day,
                            "eps_wrk_hr": eps_wrk_hr,
                            "eps_paid_stat": eps_paid_stat,
                            "pay_tax": pay_tax,
                        }

                        if (eps_inde!='S'):
                            employee_paysum_list.append(record)
                            if eps_inde=='I':
                                employee_paysum_income_list.append(record)

                            if eps_inde=='D':
                                employee_paysum_deduct_list.append(record)


                else:
                    print("Error")



            # EMP_EXEPND_LIST
            employee_expend_list = []
            employee_expend_obj = None        
            record = {}

            sql = "select a.*,b.emp_fname_th,b.emp_lname_th,b.emp_rank,c.rank_en "
            sql += ",b.emp_type,d.pay_th,e.dcp_th  from emp_expend as A "
            sql += "left join  employee as B on a.exp_emp_id=b.emp_id "
            sql += "left join  com_rank as C on b.emp_rank=c.rank_id "
            sql += "left join  t_paytype as D on a.exp_pay_type=d.pay_type "
            sql += "left join  t_discipline as E on a.exp_dcp_id=e.dcp_id "
            sql += "where (a.exp_no<>'') and (a.exp_prd_frm='" + str(period_option) + "' or a.exp_prd_id='" + str(period_option) + "') and a.exp_emp_id=" + str(emp_id) + " "
            sql += "order by a.exp_pay_type, a.exp_date, a.exp_emp_id;"
            print("DEBUG SQL : ", sql)
            try:
                cursor = connection.cursor()
                cursor.execute(sql)
                employee_expend_obj = cursor.fetchall()
            except db.OperationalError as e:
                is_error = True
                error_message = "<b>Error: please send this error to IT team</b><br>" + str(e)
            except db.Error as e:
                is_error = True
                error_message = "<b>Error: please send this error to IT team</b><br>" + str(e)
            finally:
                cursor.close()

            record = {}
            row_count = 1
            if employee_expend_obj is not None:
                if len(employee_expend_obj) > 0:                        
                    for item in employee_expend_obj:
                        exp_no = item[0]
                        exp_emp_id = item[1]
                        exp_date = item[2].strftime("%d/%m/%Y")
                        exp_pay_type = item[3]
                        exp_dcp_id = item[4]
                        exp_order = item[5]
                        exp_inde = item[6]

                        exp_amt_all = 0 if item[7] is None else '{:,}'.format(item[7])
                        exp_amt_prd = 0 if item[8] is None else '{:,}'.format(item[8])
                        exp_amt_paid = 0 if item[9] is None else '{:,}'.format(item[9])
                        exp_amt_debt = 0 if item[10] is None else '{:,}'.format(item[10])
                        exp_amt_bal = 0 if item[11] is None else '{:,}'.format(item[11])

                        exp_type = item[12]
                        exp_eff_fdate = item[13]
                        exp_eff_tdate = item[14]
                        exp_prd_frm = item[15]
                        exp_prd_to = item[16]
                        exp_prd_id = item[17]
                        pay_th = item[37]

                        record = {
                            "row_count": row_count,
                            "exp_no": exp_no,
                            "exp_emp_id": exp_emp_id,
                            "exp_no": exp_no,
                            "exp_emp_id": exp_emp_id,
                            "exp_date": exp_date,
                            "exp_pay_type": exp_pay_type,
                            "pay_th": pay_th,
                            "exp_dcp_id": exp_dcp_id,
                            "exp_order": exp_order,
                            "exp_inde": exp_inde,
                            "exp_amt_all": exp_amt_all,
                            "exp_amt_prd": exp_amt_prd,
                            "exp_amt_paid": exp_amt_paid,
                            "exp_amt_debt": exp_amt_debt,
                            "exp_amt_bal": exp_amt_bal,
                            "exp_type": exp_type,
                            "exp_eff_fdate": exp_eff_fdate.strftime("%d/%m/%Y"),
                            "exp_eff_tdate": exp_eff_tdate.strftime("%d/%m/%Y"),
                            "exp_prd_frm": exp_prd_frm,
                            "exp_prd_to": exp_prd_to,
                            "exp_prd_id": exp_prd_id,
                        }
                        employee_expend_list.append(record)
                        row_count = row_count + 1







            sql = "select prd_date_paid from t_period where prd_id='" + str(period_option) + "';"            
            paid_period_obj = None
            try:
                cursor = connection.cursor()
                cursor.execute(sql)
                paid_period_obj = cursor.fetchone()
            except db.OperationalError as e:
                is_error = True
                error_message = "<b>Error: please send this error to IT team</b><br>" + str(e)
            except db.Error as e:
                is_error = True
                error_message = "<b>Error: please send this error to IT team</b><br>" + str(e)
            finally:
                cursor.close()

            if paid_period_obj is not None:
                if len(paid_period_obj) > 0:
                    period_last_3_digits = period_option[-3:]                    
                    paid_period = paid_period_obj[0].strftime("%d/%m/%Y")
                    year_period_in_thai = str(int(paid_period[-4:]) + 543)[-2:]
                    print(period_option, period_last_3_digits, year_period_in_thai, paid_period)
            else:
                print("No record.")
        else:
            print("No record.")
    else:
        is_error = True
        error_message = "ไม่พบข้อมูล"
        emp_id = ""
        emp_title_text = ""
        emp_full_name = ""
        emp_rank = ""
        emp_status = ""
        emp_dept = ""
        dept_en = ""
        dept_en_short = ""
        emp_join_date = ""
        emp_term_date = ""
        eps_paid_stat_text = '?'
        emp_acc_no = ""
        emp_acc_bank = ""

        # Gross Income
        eps_prd_in = ""

        # Net Income
        eps_prd_net = ""

        # YTD Income
        eps_ysm_in = ""

        # YTD Prov.Func
        eps_ysm_prv = ""

        # Total Deduct
        eps_prd_de = ""

        # Tax
        eps_prd_tax = ""

        # YTD Tax
        eps_ysm_tax = ""

        # YTD Social Security
        eps_ysm_soc = ""

    document = DocxTemplate(template_name)
    style = document.styles['Normal']
    font = style.font
    font.name = 'AngsanaUPC'
    font.size = Pt(14)

    context = {
        "paid_period": paid_period,
        "period_last_3_digits":  period_last_3_digits,
        "year_period_in_thai": year_period_in_thai,    
        "emp_id": emp_id,
        "emp_title_text": emp_title_text,
        "emp_full_name": emp_full_name,
        "emp_rank": emp_rank,
        "emp_status": emp_status,
        "emp_dept": emp_dept,
        "dept_en": dept_en,
        "dept_en_short": dept_en_short,
        "emp_join_date": emp_join_date,
        "emp_term_date": emp_term_date,        
        "emp_acc_no": emp_acc_no,
        "emp_acc_bank": emp_acc_bank,
        "eps_emp_id": eps_emp_id,
        "payment_type": payment_type,
        "eps_inde": eps_inde,
        "income_or_deduct": income_or_deduct,
        "eps_comp": eps_comp,
        "eps_percent": eps_percent,
        "eps_wrk_day": eps_wrk_day,
        "eps_wrk_hr": eps_wrk_hr,
        "eps_paid_stat": eps_paid_stat,        

        "employee_paysum_list": list(employee_paysum_list),
        "employee_paysum_income_list": list(employee_paysum_income_list),
        "employee_paysum_deduct_list": list(employee_paysum_deduct_list),
        "eps_paid_stat_text": eps_paid_stat_text,
        "eps_prd_in": eps_prd_in,
        # "eps_prd_in": 1000,

        "eps_prd_net": eps_prd_net,
        "eps_ysm_in": eps_ysm_in,
        "eps_ysm_prv": eps_ysm_prv,
        "eps_prd_de": eps_prd_de,
        "eps_prd_tax": eps_prd_tax,
        "eps_ysm_tax": eps_ysm_tax,
        "eps_ysm_soc": eps_ysm_soc,
        "employee_expend_list": list(employee_expend_list),
    }

    document.render(context)
    document.save(MEDIA_ROOT + '/monitoring/download/' + file_name + ".docx")    

    # TODO
    docx_file = path.abspath("media\\monitoring\\download\\" + file_name + ".docx")
    pdf_file = path.abspath("media\\monitoring\\download\\" + file_name + ".pdf")    
    convert(docx_file, pdf_file)

    return FileResponse(open(pdf_file, 'rb'), content_type='application/pdf')


@permission_required('dailyattendreport.can_access_post_manpower_report', login_url='/accounts/login/')
def AjaxPostManpowerReport(request):    
    is_error = True
    message = "TODO"

    today_date = settings.TODAY_DATE.strftime("%Y-%m-%d")
    contract_number_from = request.POST.get('contract_number_from')
    contract_number_to = request.POST.get('contract_number_to')
    contract_zone = request.POST.get('contract_zone')
    start_date = request.POST.get('start_date')
    end_date = request.POST.get('end_date')        
    # print("DEBUG: ", today_date, contract_number_from, contract_number_to, contract_zone, start_date, end_date)

    # Convert string to date
    sd = datetime.datetime.strptime(start_date, '%d/%m/%Y')
    ed = datetime.datetime.strptime(end_date, '%d/%m/%Y')

    # Get number of dasys    
    number_of_days = abs((ed - sd).days) + 1
    print("Days : ", number_of_days)

    if number_of_days > 31: 
        response = JsonResponse(data={        
            "is_error": True,
            "message": "สามารถเลือกจำนวนวันได้ไม่เกิน 1 เดือน",
        })
        response.status_code = 200
        return response

    # Get cnt_id list
    cnt_id_list = []
    record = {}
    try:
        cursor = connection.cursor()
        for i in range(number_of_days):                        
            sql = "select distinct h.cnt_id, h.dly_date, cus.cus_name_th, cus.cus_name_en, z.zone_en, sum(case when h.absent=0 then 1 else 0 end) as total, h.dept_id "
            if (sd.strftime("%Y-%m-%d")==today_date):
                sql += "from dly_plan h "
            else:
                sql += "from his_dly_plan h "
            sql += "left join cus_contract con on h.cnt_id=con.cnt_id "
            sql += "left join customer cus on con.cus_id=cus.cus_id and con.cus_brn=cus.cus_brn "
            sql += "left join com_zone z on cus.cus_zone=z.zone_id "
            sql += "where dly_date>='" + sd.strftime("%Y-%m-%d") + "' and dly_date<='" + sd.strftime("%Y-%m-%d") + "' "            
            sql += "and h.cnt_id>=" + contract_number_from + " and h.cnt_id<=" + contract_number_to + " and h.absent=0 "

            if contract_zone != "" and contract_zone!="all_zone":            
                sql += "and dept_id=" + contract_zone + " "
            sql += "group by h.cnt_id, h.dly_date, cus.cus_name_th, cus.cus_name_en, z.zone_en, h.dept_id "
            sql += "order by h.cnt_id"
            print("SQLLL :", sql)
            cursor.execute(sql)
            obj = cursor.fetchall()
            if obj is not None:          
                for item in obj:
                    dept_name = zone_name_display_text(str(item[6]))
                    record = {
                        "cnt_id":item[0], "dly_date":item[1].day, "cus_name_th":item[2], "cus_name_en":item[3], "zone_th":item[4], "total":item[5], "dept_name":dept_name
                    }
                    cnt_id_list.append(record)

            sd += datetime.timedelta(days=1)
        is_error = False
    finally:        
        cursor.close()
    
    # Get day list
    day_list = []
    sd = datetime.datetime.strptime(start_date, '%d/%m/%Y')
    ed = datetime.datetime.strptime(end_date, '%d/%m/%Y')
    for i in range(number_of_days):
        day_list.append(sd.strftime("%d"))
        sd += datetime.timedelta(days=1)

    unique_cnt_id_list = { each['cnt_id'] : each for each in cnt_id_list }.values()
    
    response = JsonResponse(data={   
        "is_error": is_error,
        "message": message,
        "number_of_days": number_of_days,
        "day_list": day_list,
        "unique_cnt_id_list": list(unique_cnt_id_list),
        "cnt_id_list": list(cnt_id_list),
    })

    response.status_code = 200
    return response


@permission_required('dailyattendreport.can_access_gpm_422_no_of_guard_operation_by_empl_by_zone_report', login_url='/accounts/login/')
def AjaxGPM422NoOfGuardOperationByEmplByZoneReport(request, *args, **kwargs):    
    base_url = MEDIA_ROOT + '/monitoring/template/'    
    template_name = base_url + 'GPM_422.docx'
    file_name = request.user.username + "_GPM_422"

    start_date = kwargs['start_date']
    start_date = datetime.datetime.strptime(start_date, "%d/%m/%Y").date()

    end_date = kwargs['end_date']
    end_date = datetime.datetime.strptime(end_date, "%d/%m/%Y").date()    

    dept_zone = kwargs['dept_zone']

    dly_plan_obj = None
    record = {}
    dly_plan_list = []
    error_message = ""

    # ironman
    sql = "select emp_id,emp_fname_th,emp_lname_th,sch_rank,cnt_id,cus_name_th,dept_id,dept_en "
    sql += "FROM V_HDLYPLAN "
    # sql += "WHERE DLY_DATE='" + str(start_date) + "' and absent=0 "
    sql += "where (DLY_DATE>='" + str(start_date) + "' and DLY_DATE<='" + str(end_date) + "') "
    sql += " and absent=0 "

    if (int(dept_zone) > 0):
        sql += "and dept_id=" + str(dept_zone) + " "
    else:
        sql += "and dept_id=9999 "

    sql += "order by emp_id;"
    #print("SQL: ", sql)
    
    try:                
        cursor = connection.cursor()
        cursor.execute(sql)
        dly_plan_obj = cursor.fetchall()        
    except db.OperationalError as e:
        error_message = "<b>Error: please send this error to IT team</b><br>" + str(e)
    except db.Error as e:
        error_message = "<b>Error: please send this error to IT team</b><br>" + str(e)
    finally:
        cursor.close()

    # TODO
    document = DocxTemplate(template_name)
    style = document.styles['Normal']
    font = style.font
    font.name = 'AngsanaUPC'
    font.size = Pt(14)

    if dly_plan_obj is not None:
        
        temp_dept_id = None
        row_count = 1

        for item in dly_plan_obj:
            emp_id = item[0]
            emp_full_name = item[1].strip() + " " + item[2].strip()
            sch_rank = item[3]
            cnt_id = item[4]
            cus_name_th = item[5]
            dept_id = item[6]            
            dept_en = item[7].replace("('", "")

            if temp_dept_id is None:
                table = document.add_table(rows=1, cols=6, style='TableGridLight')
                a = table.cell(0, 0)
                b = table.cell(0, 5)
                c = a.merge(b)
                c.text = '%s' % (dept_id)
                c.paragraphs[0].runs[0].font.bold = True
                c.paragraphs[0].runs[0].font.size = Pt(15)

                row = table.add_row().cells
                row[0].text = "No."                
                row[1].text = "EMP ID"
                row[2].text = "NAME"
                row[3].text = "RANK"
                row[4].text = "CNT ID"
                row[5].text = "SITE NAME"

                row[0].width = Cm(0.5)
                row[1].width = Cm(2)
                row[2].width = Cm(4)
                row[3].width = Cm(1)
                row[4].width = Cm(2)
                
                '''
                table.rows[1].cells[0].paragraphs[0].runs[0].font.bold = True
                table.rows[1].cells[0].paragraphs[0].runs[0].font.size = Pt(15)
                table.rows[1].cells[1].paragraphs[0].runs[0].font.bold = True
                table.rows[1].cells[1].paragraphs[0].runs[0].font.size = Pt(15)
                table.rows[1].cells[2].paragraphs[0].runs[0].font.bold = True
                table.rows[1].cells[2].paragraphs[0].runs[0].font.size = Pt(15)
                table.rows[1].cells[3].paragraphs[0].runs[0].font.bold = True
                table.rows[1].cells[3].paragraphs[0].runs[0].font.size = Pt(15)
                table.rows[1].cells[4].paragraphs[0].runs[0].font.bold = True
                table.rows[1].cells[4].paragraphs[0].runs[0].font.size = Pt(15)
                table.rows[1].cells[5].paragraphs[0].runs[0].font.bold = True
                table.rows[1].cells[5].paragraphs[0].runs[0].font.size = Pt(15)
                '''

                if dept_id is not None:
                    row = table.add_row().cells
                    row[0].text = str(row_count)
                    row[1].text = str(emp_id)
                    row[2].text = str(emp_full_name)
                    row[3].text = str(sch_rank)
                    row[4].text = str(cnt_id)
                    row[5].text = str(cus_name_th)

                    row[0].width = Cm(0.5)
                    row[1].width = Cm(2)
                    row[2].width = Cm(4)
                    row[3].width = Cm(1)
                    row[4].width = Cm(2)

                    '''
                    table.rows[1].cells[0].paragraphs[0].runs[0].font.bold = True
                    table.rows[1].cells[0].paragraphs[0].runs[0].font.size = Pt(15)
                    table.rows[1].cells[1].paragraphs[0].runs[0].font.bold = True
                    table.rows[1].cells[1].paragraphs[0].runs[0].font.size = Pt(15)
                    table.rows[1].cells[2].paragraphs[0].runs[0].font.bold = True
                    table.rows[1].cells[2].paragraphs[0].runs[0].font.size = Pt(15)
                    table.rows[1].cells[3].paragraphs[0].runs[0].font.bold = True
                    table.rows[1].cells[3].paragraphs[0].runs[0].font.size = Pt(15)
                    table.rows[1].cells[4].paragraphs[0].runs[0].font.bold = True
                    table.rows[1].cells[4].paragraphs[0].runs[0].font.size = Pt(15)
                    table.rows[1].cells[5].paragraphs[0].runs[0].font.bold = True
                    table.rows[1].cells[5].paragraphs[0].runs[0].font.size = Pt(15)
                    '''

                    zone_name = "     " + str(dept_en)
                    row = table.rows[0]
                    zone_name = row.cells[0].paragraphs[0].add_run(zone_name)
                    zone_name.font.name = 'AngsanaUPC'
                    zone_name.font.size = Pt(16)
                    zone_name.bold = True
            else:
                if dept_id != temp_dept_id:
                    p = document.add_paragraph()
                    runner = p.add_run('TOTAL  %s' % str(row_count - 1))
                    runner.bold = True
                    zone_name.font.name = 'AngsanaUPC'
                    runner.font.size = Pt(15)

                    table = document.add_table(rows=1, cols=6, style='TableGridLight')                    

                    a = table.cell(0, 0)
                    b = table.cell(0, 1)
                    c = table.cell(0, 5)
                    d = a.merge(c)
                    d.text = '%s' % (dept_id)
                    d.paragraphs[0].runs[0].font.bold = True
                    d.paragraphs[0].runs[0].font.size = Pt(15)

                    row_count = 1                                
                    row = table.add_row().cells
                    row[0].text = str(row_count)
                    row[1].text = str(emp_id)
                    row[2].text = str(emp_full_name)
                    row[3].text = str(sch_rank)
                    row[4].text = str(cnt_id)
                    row[5].text = str(cus_name_th)

                    row[0].width = Cm(0.5)
                    row[1].width = Cm(2)
                    row[2].width = Cm(4)
                    row[3].width = Cm(1)
                    row[4].width = Cm(2)

                    '''
                    table.rows[1].cells[0].paragraphs[0].runs[0].font.bold = True
                    table.rows[1].cells[0].paragraphs[0].runs[0].font.size = Pt(15)
                    table.rows[1].cells[1].paragraphs[0].runs[0].font.bold = True
                    table.rows[1].cells[1].paragraphs[0].runs[0].font.size = Pt(15)
                    table.rows[1].cells[2].paragraphs[0].runs[0].font.bold = True
                    table.rows[1].cells[2].paragraphs[0].runs[0].font.size = Pt(15)
                    table.rows[1].cells[3].paragraphs[0].runs[0].font.bold = True
                    table.rows[1].cells[3].paragraphs[0].runs[0].font.size = Pt(15)
                    table.rows[1].cells[4].paragraphs[0].runs[0].font.bold = True
                    table.rows[1].cells[4].paragraphs[0].runs[0].font.size = Pt(15)
                    table.rows[1].cells[5].paragraphs[0].runs[0].font.bold = True
                    table.rows[1].cells[5].paragraphs[0].runs[0].font.size = Pt(15)
                    '''

                    zone_name = "     " + str(dept_en)
                    row = table.rows[0]
                    zone_name = row.cells[0].paragraphs[0].add_run(zone_name)
                    zone_name.font.name = 'AngsanaUPC'
                    zone_name.font.size = Pt(16)
                    zone_name.bold = True
                else:
                    row = table.add_row().cells
                    row[0].text = str(row_count)
                    row[1].text = str(emp_id)
                    row[2].text = str(emp_full_name)
                    row[3].text = str(sch_rank)
                    row[4].text = str(cnt_id)
                    row[5].text = str(cus_name_th)

                    row[0].width = Cm(0.5)
                    row[1].width = Cm(2)
                    row[2].width = Cm(4)
                    row[3].width = Cm(1)
                    row[4].width = Cm(2)

                    '''
                    table.rows[1].cells[0].paragraphs[0].runs[0].font.bold = True
                    table.rows[1].cells[0].paragraphs[0].runs[0].font.size = Pt(15)
                    table.rows[1].cells[1].paragraphs[0].runs[0].font.bold = True
                    table.rows[1].cells[1].paragraphs[0].runs[0].font.size = Pt(15)
                    table.rows[1].cells[2].paragraphs[0].runs[0].font.bold = True
                    table.rows[1].cells[2].paragraphs[0].runs[0].font.size = Pt(15)
                    table.rows[1].cells[3].paragraphs[0].runs[0].font.bold = True
                    table.rows[1].cells[3].paragraphs[0].runs[0].font.size = Pt(15)
                    table.rows[1].cells[4].paragraphs[0].runs[0].font.bold = True
                    table.rows[1].cells[4].paragraphs[0].runs[0].font.size = Pt(15)
                    table.rows[1].cells[5].paragraphs[0].runs[0].font.bold = True
                    table.rows[1].cells[5].paragraphs[0].runs[0].font.size = Pt(15)
                    '''

            temp_dept_id = dept_id
            row_count += 1
        
        if len(dly_plan_obj) > 0:
            p = document.add_paragraph()
            runner = p.add_run('TOTAL  %s' % str(row_count - 1))
            runner.bold = True            
            runner.font.size = Pt(15)

        context = {
            'start_date': start_date.strftime("%d/%m/%Y"),
            'end_date': start_date.strftime("%d/%m/%Y"),
            'dept_zone': dept_zone,
        }
        
        document.render(context)
        document.save(MEDIA_ROOT + '/monitoring/download/' + file_name + ".docx")        

    else:
        context = {
            'start_date': start_date.strftime("%d/%m/%Y"),
            'end_date': start_date.strftime("%d/%m/%Y"),
            'dept_zone': dept_zone,
        }
        
        document.render(context)
        document.save(MEDIA_ROOT + '/monitoring/download/' + file_name + ".docx")

    context = {
        'start_date': start_date.strftime("%d/%m/%Y"),
        'end_date': start_date.strftime("%d/%m/%Y"),
        'dept_zone': dept_zone,
    }

    document.render(context)
    document.save(MEDIA_ROOT + '/monitoring/download/' + file_name + ".docx")    

    # TODO
    docx_file = path.abspath("media\\monitoring\\download\\" + file_name + ".docx")
    pdf_file = path.abspath("media\\monitoring\\download\\" + file_name + ".pdf")    
    convert(docx_file, pdf_file)

    return FileResponse(open(pdf_file, 'rb'), content_type='application/pdf')


@permission_required('dailyattendreport.can_access_gpm403_daily_guard_performance_by_contract_report', login_url='/accounts/login/')
def AjaxGPMWorkOnDayOffReport(request, *args, **kwargs):
    print("*****************************")
    print("AjaxGPMWorkOnDayOffReport()")
    print("*****************************")
    base_url = MEDIA_ROOT + '/monitoring/template/'    
    template_name = base_url + 'GPM_HDOF.docx'
    file_name = request.user.username + "_GPM_HDOF"

    start_date = kwargs['start_date']
    start_date = datetime.datetime.strptime(start_date, "%d/%m/%Y").date()
    end_date = kwargs['end_date']
    end_date = datetime.datetime.strptime(end_date, "%d/%m/%Y").date()

    # TODO    
    sql = "SELECT emp_fname_th, emp_lname_th, shf_desc, dept_en, cnt_id, emp_id, dly_date, dept_id, sch_rank, absent, upd_by, upd_gen, cus_name_th, dof from V_HDLYPLAN "
    sql += "Where V_HDLYPLAN.dof = 1 AND V_HDLYPLAN.absent = 0 "
    sql += "and (V_HDLYPLAN.DLY_DATE>='" + str(start_date) + "' and V_HDLYPLAN.DLY_DATE<='" + str(end_date) + "') "
    sql += "order By V_HDLYPLAN.dept_id Asc"
    print("DEBUG: ", sql)

    dly_plan_obj = None
    record = {}
    dly_plan_list = []
    error_message = ""

    try:                
        cursor = connection.cursor()
        cursor.execute(sql)
        dly_plan_obj = cursor.fetchall()     
    except db.OperationalError as e:
        error_message = "<b>Error: please send this error to IT team</b><br>" + str(e)
    except db.Error as e:
        error_message = "<b>Error: please send this error to IT team</b><br>" + str(e)
    finally:
        cursor.close()

    # TODO
    document = DocxTemplate(template_name)
    style = document.styles['Normal']
    font = style.font
    font.name = 'AngsanaUPC'
    font.size = Pt(14)

    if dly_plan_obj is not None:
        
        temp_dept_id = None
        row_count = 1

        for item in dly_plan_obj:
            emp_full_name = item[0].strip() + " " + item[1].strip()            
            shf_desc = item[2].replace("('", "")            
            dept_en = item[3].replace("('", "")        

            if (dept_en.find('SP-Zone') != -1):
                dept_en_short = dept_en.replace("SP-Zone", "").lstrip()[0]
            elif (dept_en.find('ZONE H BPI') != -1):
                dept_en_short = "H"
            elif (dept_en.find('ZONE Control') != -1):
                dept_en_short = "CR"
            elif (dept_en.find('SP-Samui') != -1):
                dept_en_short = "SM"
            elif (dept_en.find('ZONE PHUKET') != -1):
                dept_en_short = "P"
            elif (dept_en.find('Zone Khon kean') != -1):
                dept_en_short = "K"
            elif (dept_en.find('BEM') != -1):
                dept_en_short = "BEM"
            elif (dept_en.find('ZONE I') != -1):
                dept_en_short = "I"
            elif (dept_en.find('ZONE R') != -1):
                dept_en_short = "R"
            elif (dept_en.find('Zone Songkhla') != -1):
                dept_en_short = "SK"
            else:
                dept_en_short = "?"

            cnt_id = item[4]
            emp_id = item[5]            
            dly_date = item[6].strftime("%d/%m/%Y")
            dept_id = item[7]
            sch_rank = item[8]
            absent = item[9]
            upd_by = item[10]
            upd_gen = item[11]
            cus_name_th = item[12]
            dof = item[13]

            if temp_dept_id is None:
                table = document.add_table(rows=1, cols=9, style='TableGridLight')                                                
                a = table.cell(0, 0)
                b = table.cell(0, 8)
                c = a.merge(b)
                c.text = '%s' % (dept_id)
                c.paragraphs[0].runs[0].font.bold = True
                c.paragraphs[0].runs[0].font.size = Pt(15)

                row = table.add_row().cells
                row[0].text = "No."
                row[1].text = "Date"
                row[2].text = "EMP ID"
                row[3].text = "Name"
                row[4].text = "Rank"
                row[5].text = "Zone"
                row[6].text = "Shift"
                row[7].text = "POST ID"
                row[8].text = "Site Name"

                row[0].width = Cm(0.5)
                row[1].width = Cm(2)
                row[2].width = Cm(2)
                row[3].width = Cm(4)
                row[4].width = Cm(1)
                row[5].width = Cm(1)
                row[6].width = Cm(5)
                row[7].width = Cm(2)
                
                if dept_id is not None:
                    row = table.add_row().cells
                    row[0].text = str(row_count)
                    row[1].text = str(dly_date)
                    row[2].text = str(emp_id)
                    row[3].text = str(emp_full_name)
                    row[4].text = str(sch_rank)
                    row[5].text = str(dept_en_short)
                    row[6].text = str(shf_desc)
                    row[7].text = str(cnt_id)
                    row[8].text = str(cus_name_th)  

                    row[0].width = Cm(0.5)
                    row[1].width = Cm(2)
                    row[2].width = Cm(2)
                    row[3].width = Cm(4)
                    row[4].width = Cm(1)
                    row[5].width = Cm(1)
                    row[6].width = Cm(5)
                    row[7].width = Cm(2)

                    company_name = "   " + str(item[3].replace("('", ""))
                    row = table.rows[0]
                    company_name = row.cells[0].paragraphs[0].add_run(company_name)
                    company_name.font.name = 'AngsanaUPC'
                    company_name.font.size = Pt(16)
                    company_name.bold = True
            else:
                if dept_id != temp_dept_id:
                    p = document.add_paragraph()
                    runner = p.add_run('TOTAL  %s' % str(row_count - 1))
                    runner.bold = True
                    company_name.font.name = 'AngsanaUPC'
                    runner.font.size = Pt(15)

                    table = document.add_table(rows=1, cols=9, style='TableGridLight')                    

                    a = table.cell(0, 0)
                    b = table.cell(0, 1)
                    c = table.cell(0, 8)
                    d = a.merge(c)
                    d.text = '%s' % (dept_id)
                    d.paragraphs[0].runs[0].font.bold = True
                    d.paragraphs[0].runs[0].font.size = Pt(15)

                    row_count = 1                                
                    row = table.add_row().cells
                    row[0].text = str(row_count)
                    row[1].text = str(dly_date)
                    row[2].text = str(emp_id)
                    row[3].text = str(emp_full_name)
                    row[4].text = str(sch_rank)
                    row[5].text = str(dept_en_short)
                    row[6].text = str(shf_desc)
                    row[7].text = str(cnt_id)
                    row[8].text = str(cus_name_th)  

                    row[0].width = Cm(0.5)
                    row[1].width = Cm(2)
                    row[2].width = Cm(2)
                    row[3].width = Cm(4)
                    row[4].width = Cm(1)
                    row[5].width = Cm(1)
                    row[6].width = Cm(5)
                    row[7].width = Cm(2)
                    
                    company_name = "   " + str(item[3].replace("('", ""))
                    row = table.rows[0]
                    company_name = row.cells[0].paragraphs[0].add_run(company_name)
                    company_name.font.name = 'AngsanaUPC'
                    company_name.font.size = Pt(16)
                    company_name.bold = True
                else:
                    row = table.add_row().cells
                    row[0].text = str(row_count)
                    row[1].text = str(dly_date)
                    row[2].text = str(emp_id)
                    row[3].text = str(emp_full_name)
                    row[4].text = str(sch_rank)
                    row[5].text = str(dept_en_short)
                    row[6].text = str(shf_desc)
                    row[7].text = str(cnt_id)
                    row[8].text = str(cus_name_th)  

                    row[0].width = Cm(0.5)
                    row[1].width = Cm(2)
                    row[2].width = Cm(2)
                    row[3].width = Cm(4)
                    row[4].width = Cm(1)
                    row[5].width = Cm(1)
                    row[6].width = Cm(5)
                    row[7].width = Cm(2)

            temp_dept_id = dept_id
            row_count += 1
    
        if len(dly_plan_obj) > 0:
            p = document.add_paragraph()
            runner = p.add_run('TOTAL  %s' % str(row_count - 1))
            runner.bold = True            
            runner.font.size = Pt(15)

        context = {
            'start_date': start_date.strftime("%d/%m/%Y"),
        }
        
        document.render(context)
        document.save(MEDIA_ROOT + '/monitoring/download/' + file_name + ".docx")        

    else:
        context = {
            'start_date': start_date.strftime("%d/%m/%Y"),
            'end_date': end_date.strftime("%d/%m/%Y"),
        }
        
        document.render(context)
        document.save(MEDIA_ROOT + '/monitoring/download/' + file_name + ".docx")


    context = {
        'start_date': start_date.strftime("%d/%m/%Y"),
    }
    document.render(context)
    document.save(MEDIA_ROOT + '/monitoring/download/' + file_name + ".docx")    

    # TODO
    docx_file = path.abspath("media\\monitoring\\download\\" + file_name + ".docx")
    pdf_file = path.abspath("media\\monitoring\\download\\" + file_name + ".pdf")    
    convert(docx_file, pdf_file)

    return FileResponse(open(pdf_file, 'rb'), content_type='application/pdf')


@permission_required('dailyattendreport.can_access_gpm403_daily_guard_performance_by_contract_report', login_url='/accounts/login/')
def GenerateGPM403DailyGuardPerformanceReport(request, *args, **kwargs):    

    print("*************************************************")
    print("GenerateGPM403DailyGuardPerformanceReport()")
    print("*************************************************")

    base_url = MEDIA_ROOT + '/monitoring/template/'
    contract_number_from = kwargs['contract_number_from']
    contract_number_to = kwargs['contract_number_to']
    start_date = kwargs['start_date']
    end_date = kwargs['end_date']

    template_name = base_url + 'GPM_403.docx'
    # template_name = base_url + 'GPM_403_P.docx'
    file_name = request.user.username + "_GPM_403"

    start_date = datetime.datetime.strptime(start_date, "%d/%m/%Y").date()
    end_date = datetime.datetime.strptime(end_date, "%d/%m/%Y").date()

    contract_number_from = 0 if contract_number_from is None else contract_number_from
    contract_number_to = 0 if contract_number_to is None else contract_number_to
    today_date = settings.TODAY_DATE.strftime("%d/%m/%Y %H:%M:%S")

    sql = "select emp_fname_th, emp_lname_th, shf_desc, dept_en, cnt_id, "
    sql += "emp_id, dly_date, sch_shift, dept_id, sch_rank, "
    sql += "absent, relieft_id, tel_man, tel_paid, ot, "
    sql += "ot_hr_amt, cus_name_th, late, late_full "
    sql += "FROM V_HDLYPLAN "
    sql += "WHERE absent = 0 AND (sch_shift <> 99 OR sch_shift <> 999) "
    sql += "and (cnt_id>=" + str(contract_number_from) + " and cnt_id<=" + str(contract_number_to) + ") "
    sql += "and (dly_date>='" + str(start_date) + "' and dly_date<='" + str(end_date) + "') "
    sql += "ORDER BY cnt_id ASC, dly_date ASC, shf_desc ASC, emp_id ASC"
    print("SQL Test: ", sql)
    
    dly_plan_obj = None
    record = {}
    dly_plan_title_list = []
    dly_plan_list = []
    error_message = ""

    try:                
        cursor = connection.cursor()
        cursor.execute(sql)
        dly_plan_obj = cursor.fetchall()        
    except db.OperationalError as e:
        error_message = "<b>Error: please send this error to IT team</b><br>" + str(e)
    except db.Error as e:
        error_message = "<b>Error: please send this error to IT team</b><br>" + str(e)
    finally:
        cursor.close()

    document = DocxTemplate(template_name)
    style = document.styles['Normal']
    font = style.font
    font.name = 'AngsanaUPC'
    font.size = Pt(14)

    if dly_plan_obj is not None:
        
        temp_cnt_id = None
        company_name = ""
        row_count = 1

        for item in dly_plan_obj:
            shf_desc = item[2].replace("('", "")
            dept_en = item[3].replace("('", "")
            cnt_id = item[4]
            emp_id = item[5]
            emp_full_name = item[0].strip() + " " + item[1].strip()
            dly_date = item[6].strftime("%d/%m/%Y")
            sch_shift = item[7]
            dept_id = item[8]
            sch_rank = item[9]
            absent = item[10]

            relieft_id = item[11] if item[11] != 0 else ""
            tel_man = "x" if item[12] else ""            
            tel_paid = "{:.2f}".format(item[13]) if item[13] > 0 else ""
            
            ot = "" if item[14] else ""
            ot_hr_amt = item[15] if item[15] else ""
            cus_name_th = item[16]
            late = "x" if item[17] else ""
            late_full = "x" if item[18] else ""

            if temp_cnt_id is None:

                '''
                table = document.add_table(rows=1, cols=13, style='TableGridLight')                                                

                a = table.cell(0, 0)
                b = table.cell(0, 12)
                c = a.merge(b)
                c.text = '%s' % (cnt_id)
                c.paragraphs[0].runs[0].font.bold = True
                c.paragraphs[0].runs[0].font.size = Pt(15)

                row = table.add_row().cells
                row[0].text = "No."
                row[0].paragraphs[0].runs[0].font.bold = True

                row[1].text = "Date"
                row[1].paragraphs[0].runs[0].font.bold = True

                row[2].text = "EMP ID"
                row[2].paragraphs[0].runs[0].font.bold = True

                row[3].text = "Name"
                row[3].paragraphs[0].runs[0].font.bold = True

                row[4].text = "Rank"
                row[4].paragraphs[0].runs[0].font.bold = True

                row[5].text = "Shift"
                row[5].paragraphs[0].runs[0].font.bold = True

                row[6].text = "Relief ID"
                row[6].paragraphs[0].runs[0].font.bold = True

                row[7].text = "OT"
                row[7].paragraphs[0].runs[0].font.bold = True

                row[8].text = "Late"
                row[8].paragraphs[0].runs[0].font.bold = True

                row[9].text = "Full"
                row[9].paragraphs[0].runs[0].font.bold = True

                row[10].text = "Amt.HR"
                row[10].paragraphs[0].runs[0].font.bold = True

                row[11].text = "Call"
                row[11].paragraphs[0].runs[0].font.bold = True

                row[12].text = "Tel Paid"
                row[12].paragraphs[0].runs[0].font.bold = True
                
                row[0].width = Cm(0.5)
                row[3].width = Cm(5)
                row[5].width = Cm(8)
                row[6].width = Cm(5)
                row[12].width = Cm(5)
                '''

                if cnt_id is not None:
                    '''
                    row = table.add_row().cells                    
                    row[0].text = str(row_count)
                    row[1].text = str(dly_date)
                    row[2].text = str(emp_id)
                    row[3].text = str(emp_full_name)
                    row[4].text = str(sch_rank)
                    row[5].text = str(shf_desc)
                    row[6].text = str(relieft_id)
                    row[7].text = str(ot)
                    row[8].text = str(late)                    
                    row[9].text = str(late_full)
                    row[10].text = str(ot_hr_amt)
                    row[11].text = str(tel_man)
                    row[12].text = str(tel_paid)

                    row[0].width = Cm(0.5)
                    row[3].width = Cm(5)
                    row[5].width = Cm(8)
                    '''
                    company_name = "  " + str(cus_name_th) + "    |    " + str(dept_id) + "   " + str(dept_en)
                    dly_plan_title_list.append(str(cnt_id) + "  " + company_name)
                    
                    '''
                    row = table.rows[0]                    
                    company_name = row.cells[0].paragraphs[0].add_run(company_name)
                    company_name.font.name = 'AngsanaUPC'
                    company_name.font.size = Pt(16)
                    company_name.bold = True
                    '''
                                    
                    record = {
                        "row_count": row_count,
                        "dly_date": dly_date,
                        "emp_id": emp_id,
                        "emp_full_name": emp_full_name,
                        "sch_rank": sch_rank,
                        "sch_shift": sch_shift,
                        "shf_desc": shf_desc,                        
                        "relieft_id": relieft_id,
                        "ot": ot,
                        "late": late,
                        "late_full": late_full,
                        "ot_hr_amt": ot_hr_amt,
                        "tel_man": tel_man,
                        "tel_paid": tel_paid,
                        "cnt_id": cnt_id,
                        "company_name": cus_name_th,
                        "dept_id": dept_id,
                        "dept_en": dept_en,
                    }

            else:
                if cnt_id != temp_cnt_id:
                    '''
                    # document.add_paragraph('TOTAL      %s' % str(row_count - 1))                                         
                    p = document.add_paragraph()
                    runner = p.add_run('TOTAL  %s' % str(row_count - 1))
                    runner.bold = True
                    # company_name.font.name = 'Cordia New (Body CS)'
                    company_name.font.name = 'AngsanaUPC'
                    runner.font.size = Pt(15)

                    # p = document.add_page_break()
                    table = document.add_table(rows=1, cols=13, style='TableGridLight')                    

                    a = table.cell(0, 0)
                    b = table.cell(0, 1)
                    c = table.cell(0, 12)
                    d = a.merge(c)
                    d.text = '%s' % (cnt_id)
                    d.paragraphs[0].runs[0].font.bold = True
                    d.paragraphs[0].runs[0].font.size = Pt(15)


                    row = table.add_row().cells

                    row[0].text = "No."
                    row[0].paragraphs[0].runs[0].font.bold = True

                    row[1].text = "Date"
                    row[1].paragraphs[0].runs[0].font.bold = True

                    row[2].text = "EMP ID"
                    row[2].paragraphs[0].runs[0].font.bold = True

                    row[3].text = "Name"
                    row[3].paragraphs[0].runs[0].font.bold = True

                    row[4].text = "Rank"
                    row[4].paragraphs[0].runs[0].font.bold = True

                    row[5].text = "Shift"
                    row[5].paragraphs[0].runs[0].font.bold = True

                    row[6].text = "Relief ID"
                    row[6].paragraphs[0].runs[0].font.bold = True

                    row[7].text = "OT"
                    row[7].paragraphs[0].runs[0].font.bold = True

                    row[8].text = "Late"
                    row[8].paragraphs[0].runs[0].font.bold = True

                    row[9].text = "Full"
                    row[9].paragraphs[0].runs[0].font.bold = True

                    row[10].text = "Amt.HR"
                    row[10].paragraphs[0].runs[0].font.bold = True

                    row[11].text = "Call"
                    row[11].paragraphs[0].runs[0].font.bold = True

                    row[12].text = "Tel Paid"
                    row[12].paragraphs[0].runs[0].font.bold = True
                    
                    row[0].width = Cm(0.5)
                    row[3].width = Cm(5)
                    row[5].width = Cm(8)
                    row[6].width = Cm(5)
                    row[12].width = Cm(5)

                    

                    row_count = 1                                
                    row = table.add_row().cells
                    row[0].text = str(row_count)
                    row[1].text = str(dly_date)
                    row[2].text = str(emp_id)
                    row[3].text = str(emp_full_name)
                    row[4].text = str(sch_rank)
                    row[5].text = str(shf_desc)
                    row[6].text = str(relieft_id)
                    row[7].text = str(ot)
                    row[8].text = str(late)                    
                    row[9].text = str(late_full)
                    row[10].text = str(ot_hr_amt)
                    row[11].text = str(tel_man)
                    row[12].text = str(tel_paid)

                    row[0].width = Cm(0.5)
                    row[3].width = Cm(5)
                    row[5].width = Cm(8)
                    '''
    
                    company_name = "  " + str(cus_name_th) + "    |    " + str(dept_id) + "   " + str(dept_en)
                    dly_plan_title_list.append(str(cnt_id) + "  " + company_name)
                    
                    '''
                    row = table.rows[0]                    
                    company_name = row.cells[0].paragraphs[0].add_run(company_name)
                    # company_name.font.name = 'Cordia New (Body CS)'
                    company_name.font.name = 'AngsanaUPC'
                    company_name.font.size = Pt(16)
                    company_name.bold = True
                    '''
                '''
                else:                   
                    row = table.add_row().cells
                    row[0].text = str(row_count)
                    row[1].text = str(dly_date)
                    row[2].text = str(emp_id)
                    row[3].text = str(emp_full_name)
                    row[4].text = str(sch_rank)
                    row[5].text = str(shf_desc)
                    row[6].text = str(relieft_id)
                    row[7].text = str(ot)
                    row[8].text = str(late)                    
                    row[9].text = str(late_full)
                    row[10].text = str(ot_hr_amt)
                    row[11].text = str(tel_man)
                    row[12].text = str(tel_paid)                    

                    row[0].width = Cm(0.5)
                    row[3].width = Cm(5)
                    row[5].width = Cm(8)
                '''
                   
                record = {
                    "row_count": row_count,
                    "dly_date": dly_date,
                    "emp_id": emp_id,
                    "emp_full_name": emp_full_name,
                    "sch_rank": sch_rank,
                    "sch_shift": sch_shift,
                    "shf_desc": shf_desc,
                    "relieft_id": relieft_id,
                    "ot": ot,
                    "late": late,
                    "late_full": late_full,
                    "ot_hr_amt": ot_hr_amt,
                    "tel_man": tel_man,
                    "tel_paid": tel_paid,
                    "cnt_id": cnt_id,
                    "company_name": cus_name_th,
                    "dept_id": dept_id,
                    "dept_en": dept_en,
                }


            dly_plan_list.append(record)

            temp_cnt_id = cnt_id
            row_count += 1
    
        if len(dly_plan_obj) > 0:
            p = document.add_paragraph()
            runner = p.add_run('TOTAL  %s' % str(row_count - 1))
            runner.bold = True  
            runner.font.size = Pt(15)

        context = {
            'start_date': start_date.strftime("%d/%m/%Y"),
            'end_date': end_date.strftime("%d/%m/%Y"),
            'dly_plan_list': list(dly_plan_list),
            'dly_plan_title_list': dly_plan_title_list,
        }
        
        document.render(context)
        document.save(MEDIA_ROOT + '/monitoring/download/' + file_name + ".docx")        

    else:
        context = {
            'start_date': start_date.strftime("%d/%m/%Y"),
            'end_date': end_date.strftime("%d/%m/%Y"),
            'dly_plan_list': list(dly_plan_list),
            'dly_plan_title_list': dly_plan_title_list,
        }
        
        document.render(context)
        document.save(MEDIA_ROOT + '/monitoring/download/' + file_name + ".docx")

    # return False

    # docx2pdf
    docx_file = path.abspath("media\\monitoring\\download\\" + file_name + ".docx")
    pdf_file = path.abspath("media\\monitoring\\download\\" + file_name + ".pdf")    
    convert(docx_file, pdf_file)

    return FileResponse(open(pdf_file, 'rb'), content_type='application/pdf')
    # return FileResponse(open(docx_file, 'rb'), content_type='application/word')



@permission_required('dailyattendreport.can_access_gpm403_daily_guard_performance_by_contract_report', login_url='/accounts/login/')
def GPM403DailyGuardPerformanceReport(request):
    page_title = settings.PROJECT_NAME
    db_server = settings.DATABASES['default']['HOST']
    project_name = settings.PROJECT_NAME
    project_version = settings.PROJECT_VERSION  
    
    today_date = settings.TODAY_DATE.strftime("%d/%m/%Y")
    contract_number_from = request.POST.get('contract_number_from')
    contract_number_to = request.POST.get('contract_number_to')
    start_date = request.POST.get('start_date')
    end_date = request.POST.get('end_date')

    contract_number_from = 0 if contract_number_from is None else contract_number_from
    contract_number_to = 9999999999 if contract_number_to is None else contract_number_to
    start_date = today_date if start_date is None else datetime.datetime.strptime(start_date, "%d/%m/%Y").date()
    end_date = today_date if end_date is None else datetime.datetime.strptime(end_date, "%d/%m/%Y").date()

    '''
    print("today_date = ", today_date)
    print("contract_number_from = ", contract_number_from)
    print("contract_number_to = ", contract_number_to)
    print("start_date = ", start_date)
    print("end_date = ", end_date)
    '''

    return render(request, 'dailyattendreport/gpm403_daily_guard_performance_by_contract_report.html',
        {
        'page_title': page_title, 
        'project_name': project_name, 
        'project_version': project_version,
        'db_server': db_server, 
        'today_date': today_date,
        'start_date': start_date,
        'end_date': end_date,
        'database': settings.DATABASES['default']['NAME'],
        'host': settings.DATABASES['default']['HOST'],
        'is_error': False,
        'dly_plan_list': None,
        })






@permission_required('dailyattendreport.can_access_gpm403_daily_guard_performance_by_contract_report', login_url='/accounts/login/')
def AjaxGPM403DailyGuardPerformanceReport(request):

    print("****************************************")
    print("AjaxGPM403DailyGuardPerformanceReport()")
    print("****************************************")

    page_title = settings.PROJECT_NAME
    db_server = settings.DATABASES['default']['HOST']
    project_name = settings.PROJECT_NAME
    project_version = settings.PROJECT_VERSION  
    
    today_date = settings.TODAY_DATE.strftime("%d/%m/%Y")
    contract_number_from = request.POST.get('contract_number_from')
    contract_number_to = request.POST.get('contract_number_to')
    start_date = request.POST.get('start_date')
    end_date = request.POST.get('end_date')

    contract_number_from = 0 if contract_number_from is None else contract_number_from
    contract_number_to = 9999999999 if contract_number_to is None else contract_number_to
    start_date = today_date if start_date is None else datetime.datetime.strptime(start_date, "%d/%m/%Y").date()
    end_date = today_date if end_date is None else datetime.datetime.strptime(end_date, "%d/%m/%Y").date()

    sql = "select emp_fname_th, emp_lname_th, shf_desc, dept_en, cnt_id, "
    sql += "emp_id, dly_date, sch_shift, dept_id, sch_rank, "
    sql += "absent, relieft_id, tel_man, tel_paid, ot, "
    sql += "ot_hr_amt, cus_name_th, late, late_full "
    sql += "FROM V_HDLYPLAN "
    sql += "WHERE absent = 0 AND (sch_shift <> 99 OR sch_shift <> 999) "
    sql += "and (cnt_id>=" + str(contract_number_from) + " and cnt_id<=" + str(contract_number_to) + ") "
    sql += "and (dly_date>='" + str(start_date) + "' and dly_date<='" + str(end_date) + "') "
    sql += "ORDER BY cnt_id ASC, dly_date ASC, shf_desc ASC, emp_id ASC"
    print(sql)
    
    dly_plan_obj = None
    record = {}
    dly_plan_list = []
    error_message = ""

    try:                
        cursor = connection.cursor()
        cursor.execute(sql)
        dly_plan_obj = cursor.fetchall()        
    except db.OperationalError as e:
        error_message = "<b>Error: please send this error to IT team</b><br>" + str(e)
    except db.Error as e:
        error_message = "<b>Error: please send this error to IT team</b><br>" + str(e)
    finally:
        cursor.close()

    if dly_plan_obj is not None:
        if len(dly_plan_obj)>0:           
            for item in dly_plan_obj:             
                record = {
                    "emp_fname_th": item[0],
                    "emp_lname_th": item[1],
                    "shf_desc": item[2],
                    "dept_en": item[3],
                    "cnt_id": item[4],
                    "emp_id": item[5],
                    "dly_date": item[6].strftime("%d/%m/%Y"),
                    "sch_shift": item[7],
                    "dept_id": item[8],
                    "sch_rank": item[9],
                    "absent": item[10],
                    "relieft_id": item[11],
                    "tel_man": item[12],
                    "tel_paid": item[13],
                    "ot": item[14],
                    "ot_hr_amt": item[15],
                    "cus_name_th": item[16],
                    "late": item[17],
                    "late_full": item[18],
                }
                dly_plan_list.append(record)

    response = JsonResponse(data={        
        "is_error": False,
        "error_message": error_message,
        "dly_plan_list": list(dly_plan_list),
    })

    response.status_code = 200
    return response




@login_required(login_url='/accounts/login/')
def export_gpm_422_no_of_guard_operation_by_empl_by_zone_to_excel(request, *args, **kwargs):
    response = HttpResponse(content_type='application/ms-excel')
    response['Content-Disposition'] = 'attachment; filename="GPM_Work_on_Day_Off.xls"'

    report_obj = []
    start_date = kwargs['start_date']
    start_date = datetime.datetime.strptime(start_date, "%d/%m/%Y").date()
    end_date = kwargs['end_date']
    end_date = datetime.datetime.strptime(end_date, "%d/%m/%Y").date()    
    dept_zone = kwargs['dept_zone']
    dept_zone_name = ""

    wb = xlwt.Workbook(encoding='utf-8')
    ws = wb.add_sheet('No. of Guard Operation')

    sql = "select emp_id,emp_fname_th,emp_lname_th,sch_rank,cnt_id,cus_name_th,cus_name_en,dept_id,dept_en,dly_date "
    sql += "FROM V_HDLYPLAN "    
    # sql += "WHERE DLY_DATE='" + str(start_date) + "' and absent=0 "    
    sql += "where (DLY_DATE>='" + str(start_date) + "' and DLY_DATE<='" + str(end_date) + "') "
    sql += " and absent=0 "

    if (int(dept_zone) > 0):
        sql += "and dept_id=" + str(dept_zone) + " "
    else:
        sql += "and dept_id=9999 "
    sql += "order by emp_id;"
    
    print("SQL GPM 422 : ", sql)

    try:  
        cursor = connection.cursor()
        cursor.execute(sql)
        report_obj = cursor.fetchall()        
    except db.OperationalError as e:
        error_message = "<b>Error: please send this error to IT team</b><br>" + str(e)
    except db.Error as e:
        error_message = "<b>Error: please send this error to IT team</b><br>" + str(e)
    finally:
        cursor.close()
    
    '''
    font_style = xlwt.XFStyle()
    font_style.font.bold = True
    font_style = xlwt.easyxf('font: bold 1, height 200;')
    ws.write(0, 0, "GPM 422 - No. of Guard Operation by Employee, Zone | " + "โซน " + str(dept_zone) + " ช่วงวันที่ " + str(start_date.strftime("%d/%m/%Y")) + " - " + str(end_date.strftime("%d/%m/%Y")), font_style)
    '''

    ws.col(0).width = int(5*260)
    ws.col(1).width = int(10*260)
    ws.col(2).width = int(10*260)
    ws.col(3).width = int(20*260)
    ws.col(4).width = int(5*260)
    ws.col(5).width = int(12*260)
    ws.col(6).width = int(35*260)
    ws.col(7).width = int(35*260)
    
    font_style = xlwt.XFStyle()
    font_style = xlwt.easyxf('font: bold 1, height 180;')
    columns = ['No', 'Dly_Date', 'EMP ID', 'Name', 'Rank', 'CNT ID', 'Site Name (TH)', 'Site Name (EN)']
    for col_num in range(len(columns)):
        ws.write(1, col_num, columns[col_num], font_style)

    font_style = xlwt.XFStyle()
    font_style = xlwt.easyxf('font: height 180;')

    row_num = 2
    counter = 1

    if len(report_obj) > 0:
        for row in report_obj:
            emp_id = row[0]
            emp_fname_th = row[1]
            emp_lname_th = row[2]
            emp_full_name = emp_fname_th.strip() + " " + emp_lname_th.strip()
            sch_rank = row[3]
            cnt_id = row[4]
            dept_id = row[7]
            dept_en = row[8]
            dept_zone_name = dept_en
            cus_name_th = row[5]
            cus_name_en = row[6]
            dly_date = row[9].strftime("%d/%m/%Y")

            for col_num in range(len(row)):
                if col_num==0:
                    ws.write(row_num, 0, counter, font_style)
                elif col_num==1:    
                    ws.write(row_num, col_num, dly_date, font_style)
                elif col_num==2:
                    ws.write(row_num, col_num, emp_id, font_style)
                elif col_num==3:
                    ws.write(row_num, col_num, emp_full_name, font_style)
                elif col_num==4:
                    ws.write(row_num, col_num, sch_rank, font_style)
                elif col_num==5:
                    ws.write(row_num, col_num, cnt_id, font_style)
                elif col_num==6:
                    ws.write(row_num, col_num, cus_name_th, font_style)
                elif col_num==7:
                    ws.write(row_num, col_num, cus_name_en, font_style)

            row_num += 1
            counter += 1
    else:
        ws.write(row_num, 0, "ไม่มีข้อมูล", font_style)

    font_style = xlwt.XFStyle()
    font_style.font.bold = True
    font_style = xlwt.easyxf('font: bold 1, height 200;')
    ws.write(0, 0, "GPM 422 - No. of Guard Operation by Employee, Zone | " + "Zone " + str(dept_zone) + " : " + str(dept_zone_name.strip()) + " Date " + str(start_date.strftime("%d/%m/%Y")) + " - " + str(end_date.strftime("%d/%m/%Y")), font_style)

    wb.save(response)
    return response


def zone_name_display_text(zone_id):
    zone_name = "N/A"
    
    if zone_id=="0":
        zone_name = "Zone F"
    elif zone_id=="2050":
        zone_name = "A"
    elif zone_id=="2051":
        zone_name = "B"
    elif zone_id=="2052":
        zone_name = "C1"
    elif zone_id=="2053":
        zone_name = "D"
    elif zone_id=="2054":
        zone_name = "E"
    elif zone_id=="2055":
        zone_name = "F"
    elif zone_id=="2056":
        zone_name = "G"
    elif zone_id=="2057":
        zone_name = "H"
    elif zone_id=="2058":
        zone_name = "S"
    elif zone_id=="2059":
        zone_name = "CR"
    elif zone_id=="2060":
        zone_name = "PID"
    elif zone_id=="2061":
        zone_name = "SM"
    elif zone_id=="2062":
        zone_name = "P"
    elif zone_id=="2063":
        zone_name = "Nakornsrithamrat"
    elif zone_id=="2064":
        zone_name = "Krabi"
    elif zone_id=="2065":
        zone_name = "Suratthani"
    elif zone_id=="2066":
        zone_name = "Udonthani"
    elif zone_id=="2067":
        zone_name = "SP-ZoneC2"
    elif zone_id=="2068":
        zone_name = "K"
    elif zone_id=="2069":
        zone_name = "C"
    elif zone_id=="2070":
        zone_name = "BEM"
    elif zone_id=="2071":
        zone_name = "I"
    elif zone_id=="2073":
        zone_name = "R"
    elif zone_id=="3099":
        zone_name = "ES Engineer"
    elif zone_id=="3335":
        zone_name = "SK"
    elif zone_id=="9999":
        zone_name = "Zone 0"

    return zone_name


@login_required(login_url='/accounts/login/')
def export_post_manpower_to_excel(request, *args, **kwargs):
    response = HttpResponse(content_type='application/ms-excel')
    response['Content-Disposition'] = 'attachment; filename="Post Manpower Report.xls"'

    pickup_record = []
    context = {}    
    
    today_date = settings.TODAY_DATE.strftime("%Y-%m-%d")
    contract_number_from = kwargs['contract_number_from']
    contract_number_to = kwargs['contract_number_to']
    contract_zone = kwargs['contract_zone_id']
    
    start_date = kwargs['contract_start_date']
    end_date = kwargs['contract_end_date']

    sd = datetime.datetime.strptime(start_date, "%d/%m/%Y").date()    
    ed = datetime.datetime.strptime(end_date, "%d/%m/%Y").date()

    print("DEBUG: ", today_date, contract_number_from, contract_number_to, contract_zone, sd, ed)
    
    # Get number of dasys    
    number_of_days = abs((ed - sd).days) + 1
    print("Days : ", number_of_days)

    if number_of_days > 31: 
        response = JsonResponse(data={        
            "is_error": True,
            "message": "สามารถเลือกจำนวนวันได้ไม่เกิน 1 เดือน",
        })
        response.status_code = 200
        return response

    # Get cnt_id list
    cnt_id_list = []
    record = {}
    try:
        cursor = connection.cursor()
        for i in range(number_of_days):                        
            sql = "select distinct h.cnt_id, h.dly_date, cus.cus_name_th, cus.cus_name_en, z.zone_en, sum(case when h.absent=0 then 1 else 0 end) as total, h.dept_id "
            if (sd.strftime("%Y-%m-%d")==today_date):
                sql += "from dly_plan h "
            else:
                sql += "from his_dly_plan h "
            sql += "left join cus_contract con on h.cnt_id=con.cnt_id "
            sql += "left join customer cus on con.cus_id=cus.cus_id and con.cus_brn=cus.cus_brn "
            sql += "left join com_zone z on cus.cus_zone=z.zone_id "
            sql += "where dly_date>='" + sd.strftime("%Y-%m-%d") + "' and dly_date<='" + sd.strftime("%Y-%m-%d") + "' "            
            sql += "and h.cnt_id>=" + contract_number_from + " and h.cnt_id<=" + contract_number_to + " and h.absent=0 "
            if contract_zone != "" and contract_zone!="all_zone":
                sql += "and h.dept_id=" + contract_zone + " "
            sql += "group by h.cnt_id, h.dly_date, cus.cus_name_th, cus.cus_name_en, z.zone_en, h.dept_id "
            sql += "order by h.cnt_id"
            # print("SQL :", sql)

            cursor.execute(sql)
            obj = cursor.fetchall()
            if obj is not None:            
                for item in obj:                    
                    record = {
                        "cnt_id":item[0], "dly_date":item[1].day, "cus_name_th":item[2], "cus_name_en":item[3], "zone_en":item[4], "total":item[5], "dept_id":item[6]
                    }
                    cnt_id_list.append(record)

            sd += datetime.timedelta(days=1)
        is_error = False
    finally:        
        cursor.close()
    
    # Get day list
    day_list = []
    sd = datetime.datetime.strptime(start_date, '%d/%m/%Y')
    ed = datetime.datetime.strptime(end_date, '%d/%m/%Y')
    for i in range(number_of_days):
        day_list.append(sd.strftime("%d"))
        sd += datetime.timedelta(days=1)

    unique_cnt_id_list = { each['cnt_id'] : each for each in cnt_id_list }.values()

    # Excel section
    wb = xlwt.Workbook(encoding='utf-8')
    ws = wb.add_sheet('Post Manpower')

    font_style = xlwt.XFStyle()
    # font_style.font.bold = True
    font_style = xlwt.easyxf('font: bold 1, height 200;')
    sd = datetime.datetime.strptime(start_date, '%d/%m/%Y')
    ed = datetime.datetime.strptime(end_date, '%d/%m/%Y')    
    ws.write(0, 0, "Period : " + str(sd.strftime("%d/%m/%Y")) + " - " + str(ed.strftime("%d/%m/%Y")), font_style)

    ws.col(0).width = int(5*260)
    ws.col(1).width = int(12*260)
    ws.col(2).width = int(40*260)
    ws.col(3).width = int(5*260)

    # Determine number of days        
    columns = ['No', 'Cnt ID', 'Customer Name', 'Zone']
    for col in day_list:        
        columns.append(col + "-" + sd.strftime("%m") + "-" + sd.strftime("%Y"))
        sd += datetime.timedelta(days=1)
        ws.col(3+int(col)).width = int(10*260)

    # Column Header
    font_style = xlwt.easyxf('font: bold 1, height 180;')
    for col_num in range(len(columns)):
        ws.write(1, col_num, columns[col_num], font_style)

    # Column Detail
    # font_style = xlwt.easyxf('height 180;')
    # font_style.font.bold = False
    font_style = xlwt.easyxf('font: bold 0, height 180;')    
    row_num = 2
    counter = 1

    if len(unique_cnt_id_list) > 0:
        for row in unique_cnt_id_list:
            cnt_id = row["cnt_id"]
            cus_name_en = row["cus_name_en"]
            dept_id = str(row["dept_id"]).strip()

            for col_num in range(len(row)):
                if col_num==0:
                    ws.write(row_num, 0, counter, font_style)
                elif col_num==1:
                    ws.write(row_num, col_num, cnt_id, font_style)
                elif col_num==2:
                    ws.write(row_num, col_num, cus_name_en, font_style)
                elif col_num==3:
                    ws.write(row_num, col_num, zone_name_display_text(dept_id), font_style)

            row_num += 1
            counter += 1

    # TODO    
    col_num = 4
    for count in range(number_of_days):
        row_num = 2
        for cnt_id in unique_cnt_id_list:
            for item in cnt_id_list:
                if int(cnt_id["cnt_id"]) == int(item["cnt_id"]):
                    if int(day_list[count])==int(item["dly_date"]):                                    
                        ws.write(row_num, col_num, item["total"], font_style)            
            row_num = row_num + 1
        col_num = col_num + 1
        row_num = 2

    wb.save(response)
    return response


@login_required(login_url='/accounts/login/')
def export_gpm_work_on_day_off_to_excel(request, *args, **kwargs):
    response = HttpResponse(content_type='application/ms-excel')
    response['Content-Disposition'] = 'attachment; filename="GPM_Work_on_Day_Off.xls"'

    r_d500_obj = []
    pickup_record = []
    context = {}

    start_date = kwargs['start_date']
    start_date = datetime.datetime.strptime(start_date, "%d/%m/%Y").date()
    end_date = kwargs['end_date']
    end_date = datetime.datetime.strptime(end_date, "%d/%m/%Y").date()

    wb = xlwt.Workbook(encoding='utf-8')
    ws = wb.add_sheet('Work on Day Off')

    sql = "select emp_fname_th, emp_lname_th, shf_desc, dept_en, cnt_id, emp_id, dly_date, dept_id, sch_rank, absent, upd_by, upd_gen, cus_name_th, cus_name_en, dof from V_HDLYPLAN "
    sql += "Where V_HDLYPLAN.dof = 1 AND V_HDLYPLAN.absent = 0 "
    sql += "and (V_HDLYPLAN.DLY_DATE>='" + str(start_date) + "' and V_HDLYPLAN.DLY_DATE<='" + str(end_date) + "') "
    sql += "order By V_HDLYPLAN.dept_id asc;"
    print("SQL : ", sql)

    try:
        cursor = connection.cursor()
        cursor.execute(sql)
        report_obj = cursor.fetchall()
    finally:
        cursor.close()

    font_style = xlwt.XFStyle()
    font_style.font.bold = True
    font_style = xlwt.easyxf('font: bold 1, height 200;')
    ws.write(0, 0, "GPM - Work on Day Off : " + str(start_date.strftime("%d/%m/%Y")) + " - " + str(end_date.strftime("%d/%m/%Y")), font_style)

    ws.col(0).width = int(5*260)
    ws.col(1).width = int(10*260)
    ws.col(2).width = int(10*260)
    ws.col(3).width = int(20*260)
    ws.col(4).width = int(5*260)
    ws.col(5).width = int(8*260)
    ws.col(6).width = int(15*260)
    ws.col(7).width = int(18*260)
    ws.col(8).width = int(12*260)
    ws.col(9).width = int(30*260)    
    ws.col(10).width = int(30*260)

    font_style = xlwt.XFStyle()    
    font_style = xlwt.easyxf('font: bold 1, height 180;')
    columns = ['No', 'Date', 'EMP ID', 'Name', 'Rank', 'Zone ID', 'Zone Name', 'Shift', 'CNT ID', 'CUS NAME (TH)', 'CUS NAME (EN)']
    for col_num in range(len(columns)):
        ws.write(1, col_num, columns[col_num], font_style)
    
    font_style = xlwt.XFStyle()
    font_style = xlwt.easyxf('font: height 180;')

    row_num = 2
    counter = 1

    if len(report_obj) > 0:
        for row in report_obj:
            emp_id = row[5]
            emp_fname_th = row[0]
            emp_lname_th = row[1]
            emp_full_name = emp_fname_th.strip() + " " + emp_lname_th.strip()
            dly_date = str(row[6].strftime("%d/%m/%Y"))
            sch_rank = row[8]
            dept_id = row[7]
            dept_en = row[3]

            if (dept_en.find('SP-Zone') != -1):
                dept_en_short = dept_en.replace("SP-Zone", "").lstrip()[0]
            elif (dept_en.find('ZONE H BPI') != -1):
                dept_en_short = "H"
            elif (dept_en.find('ZONE Control') != -1):
                dept_en_short = "CR"
            elif (dept_en.find('SP-Samui') != -1):
                dept_en_short = "SM"
            elif (dept_en.find('ZONE PHUKET') != -1):
                dept_en_short = "P"
            elif (dept_en.find('Zone Khon kean') != -1):
                dept_en_short = "K"
            elif (dept_en.find('BEM') != -1):
                dept_en_short = "BEM"
            elif (dept_en.find('ZONE I') != -1):
                dept_en_short = "I"
            elif (dept_en.find('ZONE R') != -1):
                dept_en_short = "R"
            elif (dept_en.find('Zone Songkhla') != -1):
                dept_en_short = "SK"
            else:
                dept_en_short = "?"
                            
            shf_desc = row[2]
            cnt_id = row[4]
            cus_name_th = row[12]
            cus_name_en = row[13]

            for col_num in range(len(row)):
                if col_num==0:
                    ws.write(row_num, 0, counter, font_style)
                elif col_num==1:
                    ws.write(row_num, col_num, dly_date, font_style)
                elif col_num==2:
                    ws.write(row_num, col_num, emp_id, font_style)
                elif col_num==3:
                    ws.write(row_num, col_num, emp_full_name, font_style)
                elif col_num==4:
                    ws.write(row_num, col_num, sch_rank, font_style)
                elif col_num==5:
                    ws.write(row_num, col_num, dept_id, font_style)
                elif col_num==6:
                    ws.write(row_num, col_num, dept_en_short, font_style)
                elif col_num==7:
                    ws.write(row_num, col_num, shf_desc, font_style)
                elif col_num==8:
                    ws.write(row_num, col_num, cnt_id, font_style)
                elif col_num==9:
                    ws.write(row_num, col_num, cus_name_th, font_style)
                elif col_num==10:
                    ws.write(row_num, col_num, cus_name_en, font_style)

            row_num += 1
            counter += 1
    else:
        ws.write(row_num, 0, "ไม่มีข้อมูล", font_style)

    wb.save(response)
    return response


@login_required(login_url='/accounts/login/')  
def export_gpm_403_daily_guard_performance_by_contract_to_excel(request, *args, **kwargs):
    response = HttpResponse(content_type='application/ms-excel')
    response['Content-Disposition'] = 'attachment; filename="GPM_403_Daily_Guard_Performance_by_Contract.xls"'

    gpm_403_obj = []
    pickup_record = []
    context = {}

    today_date = settings.TODAY_DATE.strftime("%d/%m/%Y")
    contract_number_from = kwargs['contract_number_from']
    contract_number_to = kwargs['contract_number_to']
    start_date = kwargs['start_date']
    start_date = datetime.datetime.strptime(start_date, "%d/%m/%Y").date()    
    end_date = kwargs['end_date']
    end_date = datetime.datetime.strptime(end_date, "%d/%m/%Y").date()
    
    wb = xlwt.Workbook(encoding='utf-8')
    ws = wb.add_sheet('GPM 403 Daily Guard Performance')
    
    font_style = xlwt.XFStyle()
    font_style.font.bold = True
    font_style = xlwt.easyxf('font: bold 1, height 200;')
    ws.write(0, 0, "GPM 403 - Daily Guard Performance by Contract : " + str(start_date.strftime("%d/%m/%Y")) + " - " + str(end_date.strftime("%d/%m/%Y")), font_style)

    font_style = xlwt.XFStyle()
    font_style = xlwt.easyxf('font: bold 1, height 180;')

    columns = ['No', 'Date', 'EMP ID', 'Name', 'Rank', 'Shift', 'CNT ID', 'CUS NAME (TH)', 'CUS NAME (EN)', 'Zone ID', 'Zone Name', 'Relief ID', 'OT', 'Late', 'Full', 'Amt.HR', 'Call', 'Tel Paid']
    for col_num in range(len(columns)):
        ws.write(1, col_num, columns[col_num], font_style)

    font_style = xlwt.XFStyle()
    font_style = xlwt.easyxf('font: height 180;')
    ws.col(0).width = int(5*260)
    ws.col(1).width = int(10*260)
    ws.col(2).width = int(8*260)
    ws.col(3).width = int(25*260)
    ws.col(4).width = int(5*260)
    ws.col(5).width = int(18*260)
    ws.col(6).width = int(12*260)
    ws.col(7).width = int(25*260)
    ws.col(8).width = int(25*260)
    ws.col(9).width = int(8*260)
    ws.col(10).width = int(15*260)
    ws.col(11).width = int(10*260)
    ws.col(12).width = int(10*260)
    ws.col(13).width = int(10*260)
    ws.col(14).width = int(10*260)

    sql = "select emp_fname_th, emp_lname_th, shf_desc, dept_en, cnt_id, "
    sql += "emp_id, dly_date, sch_shift, dept_id, sch_rank, "
    sql += "absent, relieft_id, tel_man, tel_paid, ot, "
    sql += "ot_hr_amt, cus_name_th, cus_name_en, late, late_full "
    sql += "FROM V_HDLYPLAN "
    sql += "WHERE absent = 0 AND (sch_shift <> 99 OR sch_shift <> 999) "
    sql += "and (cnt_id>=" + str(contract_number_from) + " and cnt_id<=" + str(contract_number_to) + ") "
    sql += "and (dly_date>='" + str(start_date) + "' and dly_date<='" + str(end_date) + "') "
    sql += "ORDER BY cnt_id ASC, dly_date ASC, shf_desc ASC, emp_id ASC"
    print(sql)

    try:
        cursor = connection.cursor()
        cursor.execute(sql)
        gpm_403_obj = cursor.fetchall()
    finally:
        cursor.close()


    # Sheet body, remaining rows
    # font_style = xlwt.XFStyle()
    # font_style = xlwt.easyxf('font: height 180;')

    # Sheet header, first row
    row_num = 2
    counter = 1

    if gpm_403_obj is not None:

        for row in gpm_403_obj:    
            emp_id = row[5]
            emp_fname_th = row[0]
            emp_lname_th = row[1]
            emp_full_name_th = emp_fname_th.strip() + " " + emp_lname_th.strip()            
            sch_rank = row[9]
            shf_desc = row[2]
            cnt_id = row[4]
            cus_name_th = row[16]
            cus_name_en = row[17]
            dept_id = row[8]
            dept_en = row[3]
            relieft_id = row[11]
            ot = row[14]
            late = row[18]
            late_full = row[19]
            ot_hr_amt = row[15]
            tel_man = row[12]
            tel_paid = row[13]
            dly_date = row[6].strftime("%d/%m/%Y")

            search_date_from = start_date
            search_date_to = end_date
        
            for col_num in range(len(row)):
                if(col_num==0):
                    ws.write(row_num, 0, counter, font_style)
                elif(col_num==1):
                    ws.write(row_num, col_num, dly_date, font_style)
                elif(col_num==2):
                    ws.write(row_num, col_num, emp_id, font_style)
                elif(col_num==3):
                    ws.write(row_num, col_num, emp_full_name_th, font_style)
                elif(col_num==4):
                    ws.write(row_num, col_num, sch_rank, font_style)
                elif(col_num==5):
                    ws.write(row_num, col_num, shf_desc, font_style)
                elif(col_num==6):
                    ws.write(row_num, col_num, cnt_id, font_style)
                elif(col_num==7):
                    ws.write(row_num, col_num, cus_name_th, font_style)
                elif(col_num==8):
                    ws.write(row_num, col_num, cus_name_en, font_style)
                elif(col_num==9):
                    ws.write(row_num, col_num, dept_id, font_style)
                elif(col_num==10):
                    ws.write(row_num, col_num, dept_en, font_style)
                elif(col_num==11):
                    ws.write(row_num, col_num, relieft_id, font_style)
                elif(col_num==12):
                    if not ot:
                        ot = "N"
                    else:
                        ot = "Y"                  
                    ws.write(row_num, col_num, ot, font_style)
                elif(col_num==13):
                    if not late:
                        late = "N"
                    else:
                        late = "Y"
                    ws.write(row_num, col_num, late, font_style)
                elif(col_num==14):
                    if not late_full:
                        late_full = "N"
                    else:
                        late_full = "Y"
                    ws.write(row_num, col_num, late_full, font_style)
                elif(col_num==15):
                    ws.write(row_num, col_num, ot_hr_amt, font_style)
                elif(col_num==16):
                    if not tel_man:
                        tel_man = "N"
                    else:
                        tel_man = "Y"
                    ws.write(row_num, col_num, tel_man, font_style)
                elif(col_num==17):
                    ws.write(row_num, col_num, tel_paid, font_style)

            row_num += 1
            counter += 1

    wb.save(response)
    return response


@permission_required('dailyattendreport.can_access_gpm_work_on_day_off_report', login_url='/accounts/login/')
def GPMWorkOnDayOffReport(request):
    print("***********************")
    print("GPMWorkOnDayOffReport()")
    print("***********************")

    page_title = settings.PROJECT_NAME
    db_server = settings.DATABASES['default']['HOST']
    project_name = settings.PROJECT_NAME
    project_version = settings.PROJECT_VERSION  
    
    today_date = settings.TODAY_DATE.strftime("%d/%m/%Y")
    contract_number_from = request.POST.get('contract_number_from')
    contract_number_to = request.POST.get('contract_number_to')
    start_date = request.POST.get('start_date')
    end_date = request.POST.get('end_date')

    contract_number_from = 0 if contract_number_from is None else contract_number_from
    contract_number_to = 9999999999 if contract_number_to is None else contract_number_to
    start_date = today_date if start_date is None else datetime.datetime.strptime(start_date, "%d/%m/%Y").date()
    end_date = today_date if end_date is None else datetime.datetime.strptime(end_date, "%d/%m/%Y").date()

    return render(request, 'dailyattendreport/gpm_work_on_day_off_report.html',
        {
        'page_title': page_title, 
        'project_name': project_name, 
        'project_version': project_version,
        'db_server': db_server, 
        'today_date': today_date,
        'start_date': start_date,
        'end_date': end_date,
        'database': settings.DATABASES['default']['NAME'],
        'host': settings.DATABASES['default']['HOST'],
        'is_error': False,
        'dly_plan_list': None,
        })


@permission_required('dailyattendreport.can_access_post_manpower_report', login_url='/accounts/login/')
def PostManpowerReport(request):
    page_title = settings.PROJECT_NAME
    db_server = settings.DATABASES['default']['HOST']
    project_name = settings.PROJECT_NAME
    project_version = settings.PROJECT_VERSION  
    today_date = settings.TODAY_DATE.strftime("%d/%m/%Y")    

    # Get zone
    dept_zone_obj = None
    sql = "select dept_id, dept_en from COM_DEPARTMENT where dept_zone=1;"
    try:
        cursor = connection.cursor()
        cursor.execute(sql)
        dept_zone_obj = cursor.fetchall()
    finally:
        cursor.close()
    
    return render(request, 'dailyattendreport/post_manpower.html',
        {
        'page_title': page_title, 
        'project_name': project_name,
        'project_version': project_version,
        'db_server': db_server, 
        'today_date': today_date,
        'database': settings.DATABASES['default']['NAME'],
        'host': settings.DATABASES['default']['HOST'],
        "dept_zone_obj": dept_zone_obj,
        })


@permission_required('dailyattendreport.can_access_psn_slip_d1_report', login_url='/accounts/login/')
def PSNSlipD1Report(request):
    page_title = settings.PROJECT_NAME
    db_server = settings.DATABASES['default']['HOST']
    project_name = settings.PROJECT_NAME
    project_version = settings.PROJECT_VERSION  
    
    today_date = settings.TODAY_DATE.strftime("%d/%m/%Y")
    
    employee_type_option = request.POST.get('employee_type_option')
    period_option = request.POST.get('period_option')

    employee_type_list = None
    period_list = None


    # Load Employee Type
    sql = "select com_type, type_des from com_type where com_type='D1';"
    try:
        cursor = connection.cursor()
        cursor.execute(sql)
        employee_type_list = cursor.fetchall()
    finally:
        cursor.close()

    # Load Period
    sql = "select * from t_period where emp_type='D1' and (prd_process=1) order by prd_id desc;"
    try:
        cursor = connection.cursor()
        cursor.execute(sql)
        period_list = cursor.fetchall()
    finally:
        cursor.close()

    '''
    if period_list is not None:
        if len(period_list)>0:
            prd_id = period_list[0][0]
            prd_year = period_list[0][1]
            prd_month = period_list[0][2]
            prd_date_frm = period_list[0][6]
            prd_date_to = period_list[0][7]
            prd_date_paid = period_list[0][8]
    '''

    return render(request, 'dailyattendreport/psn_slip_d1_report.html',
        {
        'page_title': page_title, 
        'project_name': project_name,
        'project_version': project_version,
        'db_server': db_server, 
        'today_date': today_date,
        'database': settings.DATABASES['default']['NAME'],
        'host': settings.DATABASES['default']['HOST'],
        'is_error': False,
        'employee_type_option': employee_type_option,
        'period_option': period_option,
        'employee_type_list': employee_type_list,
        'period_list': period_list,
        })



@permission_required('dailyattendreport.can_access_income_deduct_d1_report', login_url='/accounts/login/')
def IncomeDeductD1Report(request):
    page_title = settings.PROJECT_NAME
    db_server = settings.DATABASES['default']['HOST']
    project_name = settings.PROJECT_NAME
    project_version = settings.PROJECT_VERSION      
    today_date = settings.TODAY_DATE.strftime("%d/%m/%Y")
    page_title = settings.PROJECT_NAME
    db_server = settings.DATABASES['default']['HOST']
    project_name = settings.PROJECT_NAME
    project_version = settings.PROJECT_VERSION

    emp_id = request.POST.get("emp_id")
    pay_type_option = request.POST.get("pay_type_option")

    pay_type_list = []
    sql = "select * from t_paytype where pay_active=1;"
    try:
        cursor = connection.cursor()
        cursor.execute(sql)
        pay_type_list = cursor.fetchall()
    finally:
        cursor.close()

    return render(request, 'dailyattendreport/income_deduct_d1_report.html',
        {
        'page_title': page_title, 
        'project_name': project_name,
        'project_version': project_version,
        'db_server': db_server, 
        'today_date': today_date,
        'database': settings.DATABASES['default']['NAME'],
        'host': settings.DATABASES['default']['HOST'],
        'is_error': False,
        "pay_type_list": pay_type_list,
        "emp_id": emp_id,
        "pay_type_option": pay_type_option,
        })


@permission_required('dailyattendreport.can_access_gpm_422_no_of_guard_operation_by_empl_by_zone_report', login_url='/accounts/login/')
def GPM422NoOfGuardOperationByEmplByZoneReport(request):
    page_title = settings.PROJECT_NAME
    db_server = settings.DATABASES['default']['HOST']
    project_name = settings.PROJECT_NAME
    project_version = settings.PROJECT_VERSION  
    
    today_date = settings.TODAY_DATE.strftime("%d/%m/%Y")
    work_date = request.POST.get('work_date')
    dept_zone = request.POST.get('dept_zone')

    work_date = today_date if work_date is None else datetime.datetime.strptime(work_date, "%d/%m/%Y").date()
    dept_zone_obj = None

    sql = "select dept_id, dept_en from COM_DEPARTMENT where dept_zone=1;"
    try:
        cursor = connection.cursor()
        cursor.execute(sql)
        dept_zone_obj = cursor.fetchall()
    finally:
        cursor.close()

    return render(request, 'dailyattendreport/gpm_422_no_of_guard_operation_by_empl_by_zone_report.html',
        {
        'page_title': page_title, 
        'project_name': project_name, 
        'project_version': project_version,
        'db_server': db_server, 
        'today_date': today_date,
        'work_date': work_date,
        'dept_zone': dept_zone,
        'database': settings.DATABASES['default']['NAME'],
        'host': settings.DATABASES['default']['HOST'],
        'is_error': False,
        'dept_zone_obj': dept_zone_obj,
        })


@permission_required('dailyattendreport.can_access_income_deduct_d1_report', login_url='/accounts/login/')
def AjaxSearchIncomeDeductD1(request):    
    emp_id = request.POST.get('emp_id')        
    pay_type_option = request.POST.get('pay_type_option')
    period_option = 'D121011'
    # print(emp_id, pay_type_option)

    is_error = False
    error_message = "Error"
    result_list = []
    record = {}
    emp_fname_th = ""
    emp_lname_th = ""    
    emp_full_name = ""
    emp_rank = ""
    pay_type_list = []
    sql = "select * from t_paytype where pay_active=1;"
    try:
        cursor = connection.cursor()
        cursor.execute(sql)
        pay_type_list = cursor.fetchall()
    finally:
        cursor.close()

    sql = "select a.*,b.emp_fname_th,b.emp_lname_th,b.emp_rank,c.rank_en "
    sql += ",b.emp_type,d.pay_th,e.dcp_th  from emp_expend as A "
    sql += "left join  employee as B on a.exp_emp_id=b.emp_id "
    sql += "left join  com_rank as C on b.emp_rank=c.rank_id "
    sql += "left join  t_paytype as D on a.exp_pay_type=d.pay_type "
    sql += "left join  t_discipline as E on a.exp_dcp_id=e.dcp_id "
    sql += "where (a.exp_no <> '') " 
    #sql += "and (a.exp_prd_frm='" + str(period_option) + "' or a.exp_prd_id='" + str(period_option) + "') " 
    sql += "and a.exp_emp_id=" + str(emp_id) + " "
    sql += "and b.emp_type='D1' "
    if pay_type_option != "":
        sql += "and a.exp_pay_type='" + str(pay_type_option) + "' "
    sql += "order by a.exp_prd_id desc, a.exp_pay_type, a.exp_date, a.exp_emp_id;"
    print("SQLLLL : ", sql)




    try:
        cursor = connection.cursor()
        cursor.execute(sql)
        obj = cursor.fetchall()
    finally:
        cursor.close()

    if obj is not None:
        for item in obj:
            exp_prd_id = item[17]
            
            exp_emp_id = item[1]
            emp_fname_th = item[32].strip()
            emp_lname_th = item[33].strip()
            emp_rank = item[34]

            exp_doc_no = item[21]
            exp_doc_date = "" if item[22] is None else item[22].strftime("%d/%m/%Y")

            exp_date = item[2].strftime("%d/%m/%Y")
            exp_pay_type = item[3]
            exp_inde = item[6]
            
            # exp_amt_all = item[7]
            exp_amt_all = 0 if item[7] is None else '{:,}'.format(item[7])

            # exp_amt_period = item[8]
            exp_amt_period = 0 if item[8] is None else '{:,}'.format(item[8])

            # exp_amt_bal = item[11]
            exp_amt_bal = 0 if item[11] is None else '{:,}'.format(item[11])

            # exp_eff_fdate = item[13]
            exp_eff_fdate = "" if item[13] is None else item[13].strftime("%d/%m/%Y")

            # exp_eff_tdate = item[14]
            exp_eff_tdate = "" if item[14] is None else item[14].strftime("%d/%m/%Y")

            pay_th = item[37]

            # tanos
            record = {
                "exp_prd_id": exp_prd_id,
                "exp_emp_id": exp_emp_id,
                "exp_date": exp_date,
                "exp_pay_type": exp_pay_type,
                "exp_inde": exp_inde,
                "exp_amt_all": exp_amt_all,            
                "exp_amt_period": exp_amt_period,
                "exp_amt_bal": exp_amt_bal,
                "exp_eff_fdate": exp_eff_fdate,
                "exp_eff_tdate": exp_eff_tdate,
                "pay_th": pay_th,
                "exp_doc_no": exp_doc_no,
                "exp_doc_date": exp_doc_date
            }
            result_list.append(record)

    response = JsonResponse(data={   
        "is_error": is_error,
        "error_message": error_message,
        "emp_id": emp_id,
        "emp_full_name": emp_fname_th + " " + emp_lname_th,
        "emp_rank": emp_rank,
        "pay_type_option": pay_type_option,
        "pay_type_list": list(pay_type_list),
        "result_list": list(result_list),
    })

    response.status_code = 200
    return response    


@permission_required('dailyattendreport.can_access_psn_slip_d1_report', login_url='/accounts/login/')
def AjaxValidatePSNSlipD1Period(request):

    # Force to use D1
    emp_type = request.POST.get('emp_type')
    if emp_type != 'D1':
        emp_type = 'D1'

    period_option = request.POST.get('period_option')
    emp_id = request.POST.get('emp_id')    
    
    # print(emp_type, period_option, emp_id)

    if (emp_type=='' or emp_id=='' or period_option==''):
        response = JsonResponse(data={        
            "is_error": True,
            "error_message": "Please check your input data.",
        })
        response.status_code = 200
        return response

    if (emp_type != 'D1'):
        response = JsonResponse(data={        
            "is_error": True,
            "error_message": "You have not authorized to do this. This will be reported to IT.",
        })
        response.status_code = 200
        return response

    sql = "select a.*,b.dept_en,c.sts_en from employee as a left join com_department as b on a.emp_dept=b.dept_id "
    sql += "left join t_empsts as c on a.emp_status=c.sts_id where a.emp_id=" + str(emp_id) + " and a.emp_type='D1';"    
    print("SQL 111: ", sql)

    employee_info = None
    employee_paysum_list = []
    employee_expend_list = []
    eps_paid_stat_text = '?'
    emp_id = ""
    emp_full_name = ""
    emp_rank = ""
    emp_status = ""
    emp_dept = ""
    dept_en = ""
    dept_en_short = ""
    emp_join_date = ""
    emp_term_date = ""

    eps_paid_stat_text = '?'
    # Gross Income
    eps_prd_in = ""
    # Net Income
    eps_prd_net = ""
    # YTD Income
    eps_ysm_in = ""
    # YTD Prov.Func
    eps_ysm_prv = ""
    # Total Deduct
    eps_prd_de = ""
    # Tax
    eps_prd_tax = ""
    # YTD Tax
    eps_ysm_tax = ""
    # YTD Social Security
    eps_ysm_soc = ""    

    try:                
        cursor = connection.cursor()
        cursor.execute(sql)
        employee_info = cursor.fetchone()
    except db.OperationalError as e:
        is_error = True
        error_message = "<b>Error: please send this error to IT team</b><br>" + str(e)
    except db.Error as e:
        is_error = True
        error_message = "<b>Error: please send this error to IT team</b><br>" + str(e)
    finally:
        cursor.close()

    if employee_info is not None:
        is_error = False
        error_message = ""

        emp_id = employee_info[0]
        emp_full_name = str(employee_info[4]) + "  " + str(employee_info[5])
        emp_rank = employee_info[31]
        emp_status = employee_info[69]
        emp_dept = employee_info[28]
        dept_en = employee_info[68]
        if (dept_en.find('SP-Zone') != -1):
            dept_en_short = dept_en.replace("SP-Zone", "").lstrip()[0]
        elif (dept_en.find('ZONE H BPI') != -1):
            dept_en_short = "H"
        elif (dept_en.find('ZONE Control') != -1):
            dept_en_short = "CR"
        elif (dept_en.find('SP-Samui') != -1):
            dept_en_short = "SM"
        elif (dept_en.find('ZONE PHUKET') != -1):
            dept_en_short = "P"
        elif (dept_en.find('Zone Khon kean') != -1):
            dept_en_short = "K"
        elif (dept_en.find('BEM') != -1):
            dept_en_short = "BEM"
        elif (dept_en.find('ZONE I') != -1):
            dept_en_short = "I"
        elif (dept_en.find('ZONE R') != -1):
            dept_en_short = "R"
        elif (dept_en.find('Zone Songkhla') != -1):
            dept_en_short = "SK"
        else:
            dept_en_short = dept_en

        emp_join_date = "" if employee_info[39] is None else employee_info[39].strftime("%d/%m/%Y")
        emp_term_date = "" if employee_info[40] is None else employee_info[40].strftime("%d/%m/%Y")


        # TODO
        sql = "select distinct eps_prd_id from pay_sum where eps_prd_id='" + str(period_option) + "';"
        print("SQL 0: ", sql)
        try:
            cursor = connection.cursor()
            cursor.execute(sql)
            period_obj = cursor.fetchall()
        except db.OperationalError as e:
            is_error = True
            error_message = "<b>Error: please send this error to IT team</b><br>" + str(e)
        except db.Error as e:
            is_error = True
            error_message = "<b>Error: please send this error to IT team</b><br>" + str(e)
        finally:
            cursor.close()
        table = "PAY_SUM"
        if period_obj is not None:
            if len(period_obj) > 0:
                # แสดงว่าเป็นงวดปัจจุบัน
                table = "PAY_SUM"
            else:
                # แสดงว่าเป็นงวดย้อนหลัง
                table = "HIS_PAY_SUM"

        # print("TABLE = ", table)
        # Get PAYSUM
        # sql = "SELECT  a.*,b.pay_th FROM HIS_PAY_SUM as A left join t_paytype as B on a.eps_pay_type=b.pay_type "
        # sql = "SELECT  a.*,b.pay_th FROM PAY_SUM as A left join t_paytype as B on a.eps_pay_type=b.pay_type "
        sql = "SELECT a.*,b.pay_th,b.pay_tax FROM " + str(table) + " as A left join t_paytype as B on a.eps_pay_type=b.pay_type "
        sql += "where eps_prd_id='" + str(period_option) + "' and eps_emp_id=" + str(emp_id) + " "
        # sql += "and eps_inde in ('I','D') "
        sql += "ORDER BY eps_pay_type;"
        print("SQL 222: ", sql)

        employee_paysum_obj = None        
        record = {}
        try:
            cursor = connection.cursor()
            cursor.execute(sql)
            employee_paysum_obj = cursor.fetchall()
        except db.OperationalError as e:
            is_error = True
            error_message = "<b>Error: please send this error to IT team</b><br>" + str(e)
        except db.Error as e:
            is_error = True
            error_message = "<b>Error: please send this error to IT team</b><br>" + str(e)
        finally:
            cursor.close()

        if employee_paysum_obj is not None:
            if len(employee_paysum_obj) > 0:
                
                print("AAA")
                row_count = 1
                for item in employee_paysum_obj:
                    if (row_count == 1):
                        row_count = row_count + 1                        
                        eps_paid_stat = item[31]
                        if eps_paid_stat=='P':
                            eps_paid_stat_text = 'P : PAID'
                        elif eps_paid_stat=='H':
                            eps_paid_stat_text = 'H : HOLDING'
                        elif eps_paid_stat=='C':
                            eps_paid_stat_text = 'C : CHEQUE'
                        else:
                            eps_paid_stat_text = eps_paid_stat

                        # Gross Income
                        eps_prd_in = '{:,}'.format(item[27])                        

                        # Net Income
                        eps_prd_net = '{:,}'.format(item[29])

                        # YTD Income
                        eps_ysm_in = '{:,}'.format(item[14])

                        # YTD Prov.Func
                        eps_ysm_prv = '{:,}'.format(item[19])

                        # Total Deduct
                        eps_prd_de = '{:,}'.format(item[28])

                        # Tax
                        eps_prd_tax = '{:,}'.format(item[30])

                        # YTD Tax
                        eps_ysm_tax = '{:,}'.format(item[21])

                        # YTD Social Security
                        eps_ysm_soc = '{:,}'.format(item[20])

                    eps_emp_id = item[0]
                    eps_pay_type = item[2]
                    pay_th = item[37]
                    payment_type = str(eps_pay_type) + " " + str(pay_th)

                    # PAY TAX
                    # pay_tax = item[38]
                    if (item[38] == 1):
                        pay_tax = 1
                    else:
                        pay_tax = 0
                    print("pay_tax : ", pay_tax)

                    # income_or_deduct = item[7]
                    income_or_deduct = '{:,}'.format(item[7])

                    eps_inde = item[4]                    
                    
                    if (item[9] is not None):
                        eps_comp = item[9]
                    else:
                        eps_comp = ""

                    if (item[10] is not None):
                        eps_percent = item[10]
                    else:
                        eps_percent = ""

                    eps_wrk_day = item[12]
                    if (item[12] is not None):
                        eps_wrk_day = item[12]
                    else:
                        eps_wrk_day = ""

                    eps_wrk_hr = item[13]
                    if (item[13] is not None):
                        eps_wrk_hr = item[13]
                    else:
                        eps_wrk_hr = ""

                    record = {
                        "eps_emp_id": eps_emp_id,
                        "payment_type": payment_type,
                        "eps_inde": eps_inde,
                        "income_or_deduct": income_or_deduct,
                        "eps_comp": eps_comp,
                        "eps_percent": eps_percent,
                        "eps_wrk_day": eps_wrk_day,
                        "eps_wrk_hr": eps_wrk_hr,
                        "eps_paid_stat": eps_paid_stat,
                        "pay_tax": pay_tax,
                    }

                    if (eps_inde!='S'):
                        employee_paysum_list.append(record)


                # EMP_EXPEND_LIST
                # sql = "select * from emp_expend where exp_emp_id=" + str(emp_id) + " and exp_prd_id='" + str(period_option) + "' order by exp_order";
                sql = "select a.*,b.emp_fname_th,b.emp_lname_th,b.emp_rank,c.rank_en "
                sql += ",b.emp_type,d.pay_th,e.dcp_th  from emp_expend as A "
                sql += "left join  employee as B on a.exp_emp_id=b.emp_id "
                sql += "left join  com_rank as C on b.emp_rank=c.rank_id "
                sql += "left join  t_paytype as D on a.exp_pay_type=d.pay_type "
                sql += "left join  t_discipline as E on a.exp_dcp_id=e.dcp_id "
                sql += "where (a.exp_no<>'') and (a.exp_prd_frm='" + str(period_option) + "' or a.exp_prd_id='" + str(period_option) + "') and a.exp_emp_id=" + str(emp_id) + " "
                sql += "order by a.exp_pay_type, a.exp_date, a.exp_emp_id;"
                print("AAABBB : ", sql)

                employee_expend_obj = None        
                record = {}
                try:
                    cursor = connection.cursor()
                    cursor.execute(sql)
                    employee_expend_obj = cursor.fetchall()
                except db.OperationalError as e:
                    is_error = True
                    error_message = "<b>Error: please send this error to IT team</b><br>" + str(e)
                except db.Error as e:
                    is_error = True
                    error_message = "<b>Error: please send this error to IT team</b><br>" + str(e)
                finally:
                    cursor.close()

                record = {}
                if employee_expend_obj is not None:
                    if len(employee_expend_obj) > 0:                        
                        for item in employee_expend_obj:

                            exp_no = item[0]
                            exp_emp_id = item[1]
                            exp_date = item[2].strftime("%d/%m/%Y")
                            exp_pay_type = item[3]
                            exp_dcp_id = item[4]
                            exp_order = item[5]
                            exp_inde = item[6]
                            
                            # exp_amt_all = '{:,}'.format(item[7])
                            exp_amt_all = 0 if item[7] is None else '{:,}'.format(item[7])

                            # exp_amt_prd = '{:,}'.format(item[8])
                            exp_amt_prd = 0 if item[8] is None else '{:,}'.format(item[8])

                            # exp_amt_paid = '{:,}'.format(item[9])
                            exp_amt_paid = 0 if item[9] is None else '{:,}'.format(item[9])

                            # exp_amt_debt = '{:,}'.format(item[10])
                            exp_amt_debt = 0 if item[10] is None else '{:,}'.format(item[10])

                            # exp_amt_bal = '{:,}'.format(item[11])
                            exp_amt_bal = 0 if item[11] is None else '{:,}'.format(item[11])

                            exp_type = item[12]
                            exp_eff_fdate = item[13]
                            exp_eff_tdate = item[14]
                            exp_prd_frm = item[15]
                            exp_prd_to = item[16]
                            exp_prd_id = item[17]
                            pay_th = item[37]

                            record = {
                                "exp_no": exp_no,
                                "exp_emp_id": exp_emp_id,
                                "exp_no": exp_no,
                                "exp_emp_id": exp_emp_id,
                                "exp_date": exp_date,
                                "exp_pay_type": exp_pay_type,
                                "pay_th": pay_th,
                                "exp_dcp_id": exp_dcp_id,
                                "exp_order": exp_order,
                                "exp_inde": exp_inde,
                                "exp_amt_all": exp_amt_all,
                                "exp_amt_prd": exp_amt_prd,
                                "exp_amt_paid": exp_amt_paid,
                                "exp_amt_debt": exp_amt_debt,
                                "exp_amt_bal": exp_amt_bal,
                                "exp_type": exp_type,
                                "exp_eff_fdate": exp_eff_fdate.strftime("%d/%m/%Y"),
                                "exp_eff_tdate": exp_eff_tdate.strftime("%d/%m/%Y"),
                                "exp_prd_frm": exp_prd_frm,
                                "exp_prd_to": exp_prd_to,
                                "exp_prd_id": exp_prd_id,
                            }
                            employee_expend_list.append(record)

            else:
                print("No record.")
        else:
            print("No record.")
    else:
        print("BBB")
        is_error = True
        error_message = "ไม่พบข้อมูล"
        emp_id = ""
        emp_full_name = ""
        emp_rank = ""
        emp_status = ""
        emp_dept = ""
        dept_en = ""
        dept_en_short = ""
        emp_join_date = ""
        emp_term_date = ""
        eps_paid_stat_text = '?'
        # Gross Income
        eps_prd_in = ""

        # Net Income
        eps_prd_net = ""

        # YTD Income
        eps_ysm_in = ""

        # YTD Prov.Func
        eps_ysm_prv = ""

        # Total Deduct
        eps_prd_de = ""

        # Tax
        eps_prd_tax = ""

        # YTD Tax
        eps_ysm_tax = ""

        # YTD Social Security
        eps_ysm_soc = ""

    response = JsonResponse(data={   
        "is_error": is_error,
        "error_message": error_message,
        "emp_id": emp_id,
        "emp_full_name": emp_full_name,
        "emp_rank": emp_rank,
        "emp_status": emp_status,
        "emp_dept": emp_dept,
        "dept_en": dept_en,
        "dept_en_short": dept_en_short,
        "emp_join_date": emp_join_date,
        "emp_term_date": emp_term_date,
        "employee_paysum_list": list(employee_paysum_list),
        "employee_expend_list": list(employee_expend_list),
        "eps_paid_stat_text": eps_paid_stat_text,
        "eps_prd_in": eps_prd_in,
        "eps_prd_net": eps_prd_net,
        "eps_ysm_in": eps_ysm_in,
        "eps_ysm_prv": eps_ysm_prv,
        "eps_prd_de": eps_prd_de,
        "eps_prd_tax": eps_prd_tax,
        "eps_ysm_tax": eps_ysm_tax,
        "eps_ysm_soc": eps_ysm_soc,        
    })

    response.status_code = 200
    return response



@permission_required('dailyattendreport.can_access_terminate_employee_list', login_url='/accounts/login/')
def TerminateEmployeeListReport(request):
    page_title = settings.PROJECT_NAME
    db_server = settings.DATABASES['default']['HOST']
    project_name = settings.PROJECT_NAME
    project_version = settings.PROJECT_VERSION  
    
    today_date = settings.TODAY_DATE.strftime("%d/%m/%Y")
    contract_number_from = request.POST.get('contract_number_from')
    contract_number_to = request.POST.get('contract_number_to')
    start_date = request.POST.get('start_date')
    end_date = request.POST.get('end_date')

    start_date = today_date if start_date is None else datetime.datetime.strptime(start_date, "%d/%m/%Y").date()
    end_date = today_date if end_date is None else datetime.datetime.strptime(end_date, "%d/%m/%Y").date()

    dept_zone_obj = None
    sql = "select dept_id, dept_en from COM_DEPARTMENT where dept_zone=1;"
    try:
        cursor = connection.cursor()
        cursor.execute(sql)
        dept_zone_obj = cursor.fetchall()
    finally:
        cursor.close()
    
    return render(request, 'dailyattendreport/terminate_employee_list.html',
        {
        'page_title': page_title, 
        'project_name': project_name,
        'project_version': project_version,
        'db_server': db_server, 
        'today_date': today_date,
        'database': settings.DATABASES['default']['NAME'],
        'host': settings.DATABASES['default']['HOST'],
        "dept_zone_obj": dept_zone_obj,
        })


@permission_required('dailyattendreport.can_access_terminate_employee_list', login_url='/accounts/login/')
def AjaxTerminateEmployeeListReport(request):    
    employee_obj = None
    employee_list = []
    record = {}
    is_error = True
    error_message = ""

    emp_id_from = request.POST.get('emp_id_from')
    emp_id_to = request.POST.get('emp_id_to')
    
    emp_type = request.POST.get('emp_type')
    emp_type = 'D1'

    emp_dept = request.POST.get('emp_dept')
    start_date = request.POST.get('start_date')
    end_date = request.POST.get('end_date')
    
    sd = datetime.datetime.strptime(start_date, '%d/%m/%Y') 
    ed = datetime.datetime.strptime(end_date, '%d/%m/%Y') 
    # print(emp_id_from,emp_id_to,emp_type,emp_dept,sd,ed)

    sql = "select * from v_emptrm where emp_type='" + str(emp_type) + "' "
    sql += " and emp_id>='" + str(emp_id_from) + "' "
    sql += " and emp_id<='" + str(emp_id_to) + "' "
    if emp_dept!="":
        sql += " and emp_sect=" + str(emp_dept) + " "
    sql += " and emp_term_date>='" + str(sd) + "' and emp_term_date<='" + str(ed) + "';"
    print("SQL : ", sql)

    try:
        cursor = connection.cursor()
        cursor.execute(sql)
        employee_obj = cursor.fetchall()
    finally:
        cursor.close()

    if employee_obj is not None:
        for item in employee_obj:
            emp_id = item[0];
            emp_fname_th = item[1];
            emp_lname_th = item[2];
            emp_fullname_th = emp_fname_th.strip() + emp_lname_th.strip()
            emp_sect = item[5]
            emp_rank = item[6];
            emp_join_date = item[7].strftime("%d/%m/%Y");
            emp_term_date = item[8].strftime("%d/%m/%Y");
            trm_res_th = item[10];
            trm_doc_date = item[11].strftime("%d/%m/%Y");
            emp_status = item[12];
            sts_th = item[13];
            wmonth = item[21];

            if emp_status==9:
                sts_th = "ไล่ออก/ปลดออก"

            record = {
                'emp_id': emp_id,
                'emp_fname_th': emp_fname_th,
                'emp_lname_th': emp_lname_th,
                'emp_fullname_th': emp_fullname_th,
                'emp_sect': emp_sect,
                'emp_rank': emp_rank,
                'emp_join_date': emp_join_date,
                'emp_term_date': emp_term_date,
                'trm_res_th': trm_res_th,
                'trm_doc_date': trm_doc_date,
                'sts_th': sts_th,
                'wmonth': wmonth,
                'emp_status': emp_status,
            }

            employee_list.append(record)

        is_error = False
        error_message = ""

    response = JsonResponse(data={   
        "is_error": is_error,
        "error_message": error_message,
        "employee_list": employee_list,
    })

    response.status_code = 200
    return response 


@permission_required('dailyattendreport.can_access_terminate_employee_list', login_url='/accounts/login/')
def AjaxPrintTerminateEmployeeListReport(request, *args, **kwargs):    
    
    base_url = MEDIA_ROOT + '/monitoring/template/'

    emp_id_from = kwargs['emp_id_from']
    emp_id_to = kwargs['emp_id_to']
    
    emp_type = kwargs['emp_type']
    emp_type = 'D1'

    emp_dept = kwargs['emp_dept']
    start_date = kwargs['start_date']
    end_date = kwargs['end_date']
    
    sd = datetime.datetime.strptime(start_date, '%d/%m/%Y') 
    ed = datetime.datetime.strptime(end_date, '%d/%m/%Y') 

    template_name = base_url + 'Terminate_Employee_List.docx'
    file_name = request.user.username + "Terminate_Employee_List"

    sql = "select * from v_emptrm "
    sql += " where emp_type='" + str(emp_type) + "' "
    sql += " and emp_id>='" + str(emp_id_from) + "' "
    sql += " and emp_id<='" + str(emp_id_to) + "' "
    if emp_dept!="0":
        sql += " and emp_sect=" + str(emp_dept) + " "
    sql += " and emp_term_date>='" + str(sd) + "' and emp_term_date<='" + str(ed) + "';"
    
    # print(sql)
    # print(emp_id_from, emp_id_to, emp_type, emp_dept, sd, ed)
    
    employee_obj = None
    record = {}
    employee_list = []
    error_message = ""

    try:                
        cursor = connection.cursor()
        cursor.execute(sql)
        employee_obj = cursor.fetchall()        
    except db.OperationalError as e:
        error_message = "<b>Error: please send this error to IT team</b><br>" + str(e)
    except db.Error as e:
        error_message = "<b>Error: please send this error to IT team</b><br>" + str(e)
    finally:
        cursor.close()

    document = DocxTemplate(template_name)
    style = document.styles['Normal']
    font = style.font
    font.name = 'AngsanaUPC'
    font.size = Pt(14)


    if employee_obj is not None:
    
        row_count = 1

        for item in employee_obj:
            emp_id = item[0]
            emp_fname_th = item[1].strip()
            emp_lname_th = item[2].strip()
            emp_type = item[3]
            emp_com = item[4]
            emp_sect = item[5]
            emp_rank = item[6]
            
            emp_join_date = item[7].strftime('%d/%m/%Y')
            emp_term_date = item[8].strftime('%d/%m/%Y')

            title_th = item[9]
            trm_res_th = item[10]
            trm_doc_date = item[11].strftime('%d/%m/%Y')
            emp_status = item[12]
            sts_th = item[13]
            emp_fname_en = item[14]
            emp_lname_en = item[15]
            title_en = item[16]
            trm_res_en = item[17]
            emp_birth = item[18].strftime('%d/%m/%Y')
            emp_sex = item[19]
            dept_th = item[20]
            wmonth = item[21]

            record = {
                'seq': row_count,
                'emp_id': emp_id,
                'emp_fname_th': emp_fname_th,
                'emp_lname_th': emp_lname_th,
                'emp_fullname_th': title_th + " " + emp_fname_th + "  " + emp_lname_th,
                'emp_type': emp_type,
                'emp_com': emp_com,
                'emp_sect': emp_sect,
                'emp_rank': emp_rank,
                'emp_join_date': emp_join_date,
                'emp_term_date': emp_term_date,
                'title_th': title_th,
                'trm_res_th': trm_res_th,
                'trm_doc_date': trm_doc_date,
                'emp_status': emp_status,
                'sts_th': sts_th,
                'emp_fname_en': emp_fname_en,
                'emp_lname_en': emp_lname_en,
                'emp_fullname_en': title_en + " " + emp_fname_en + "  " + emp_lname_en,
                'title_en': title_en,
                'trm_res_en': trm_res_en,
                'emp_dob': emp_birth,
                'emp_sex': emp_sex,
                'dept_th': dept_th,
                'wmonth': wmonth,
            }

            employee_list.append(record)
            row_count = row_count + 1

        today_date = settings.TODAY_DATE.strftime("%d/%m/%Y %H:%M:%S")

        context = {
            'emp_type': emp_type,
            'start_date': sd.strftime("%d/%m/%Y"),
            'end_date': ed.strftime("%d/%m/%Y"),
            'employee_list': list(employee_list),
            'total_count': row_count - 1,
            'today_date': today_date,
        }
        
        document.render(context)
        document.save(MEDIA_ROOT + '/monitoring/download/' + file_name + ".docx")        

    else:
        context = {
            'start_date': start_date.strftime("%d/%m/%Y"),
            'end_date': end_date.strftime("%d/%m/%Y"),
        }
        
        document.render(context)
        document.save(MEDIA_ROOT + '/monitoring/download/' + file_name + ".docx")

    # docx2pdf
    docx_file = path.abspath("media\\monitoring\\download\\" + file_name + ".docx")
    pdf_file = path.abspath("media\\monitoring\\download\\" + file_name + ".pdf")    
    convert(docx_file, pdf_file)

    return FileResponse(open(pdf_file, 'rb'), content_type='application/pdf')



@permission_required('dailyattendreport.can_access_terminate_employee_list', login_url='/accounts/login/')
def AjaxExportTerminateEmployeeListReport(request, *args, **kwargs):
    response = HttpResponse(content_type='application/ms-excel')
    response['Content-Disposition'] = 'attachment; filename="Terminate_Employee_List.xls"'

    emp_id_from = kwargs['emp_id_from']
    emp_id_to = kwargs['emp_id_to']

    emp_type = kwargs['emp_type']
    emp_type = 'D1'

    emp_dept = kwargs['emp_dept']
    start_date = kwargs['start_date']
    end_date = kwargs['end_date']    
    sd = datetime.datetime.strptime(start_date, '%d/%m/%Y') 
    ed = datetime.datetime.strptime(end_date, '%d/%m/%Y') 
    
    employee_obj = []
    pickup_record = []
    context = {}

    wb = xlwt.Workbook(encoding='utf-8')
    ws = wb.add_sheet('Terminate Employee List')
    
    font_style = xlwt.XFStyle()
    font_style.font.bold = True
    font_style = xlwt.easyxf('font: bold 1, height 200;')
    ws.write(0, 0, "Terminate Employee List : " + str(sd.strftime("%d/%m/%Y")) + " - " + str(ed.strftime("%d/%m/%Y")), font_style)

    font_style = xlwt.XFStyle()
    font_style = xlwt.easyxf('font: bold 1, height 180;')

    columns = ['Seq', 'EMP ID', 'Name (TH)', 'Name (EN)', 'Section', 'Rank', 'Doc Date', 'Join', 'Term', 'Work', 'Status', 'Reason']
    for col_num in range(len(columns)):
        ws.write(1, col_num, columns[col_num], font_style)

    font_style = xlwt.XFStyle()
    font_style = xlwt.easyxf('font: height 180;')
    ws.col(0).width = int(5*260)
    ws.col(1).width = int(10*260)
    ws.col(2).width = int(8*260)
    ws.col(3).width = int(25*260)
    ws.col(4).width = int(5*260)
    ws.col(5).width = int(18*260)
    ws.col(6).width = int(12*260)
    ws.col(7).width = int(25*260)
    ws.col(8).width = int(25*260)
    ws.col(9).width = int(8*260)
    ws.col(10).width = int(15*260)
    ws.col(11).width = int(10*260)

    sql = "select * from v_emptrm "
    sql += " where emp_type='" + str(emp_type) + "' "
    sql += " and emp_id>='" + str(emp_id_from) + "' "
    sql += " and emp_id<='" + str(emp_id_to) + "' "
    if emp_dept!="0":
        sql += " and emp_sect=" + str(emp_dept) + " "
    sql += " and emp_term_date>='" + str(sd) + "' and emp_term_date<='" + str(ed) + "';"

    print(sql)

    try:
        cursor = connection.cursor()
        cursor.execute(sql)
        employee_obj = cursor.fetchall()
    finally:
        cursor.close()

    # Sheet header, first row
    row_num = 2
    counter = 1

    if employee_obj is not None:

        for item in employee_obj:    
            emp_id = item[0]
            emp_fname_th = item[1].strip()
            emp_lname_th = item[2].strip()
            emp_type = item[3]
            emp_com = item[4]
            emp_sect = item[5]
            emp_rank = item[6]
            
            emp_join_date = item[7].strftime('%d/%m/%Y')
            emp_term_date = item[8].strftime('%d/%m/%Y')

            title_th = item[9]
            trm_res_th = item[10]
            trm_doc_date = item[11].strftime('%d/%m/%Y')
            emp_status = item[12]
            sts_th = item[13]
            emp_fname_en = item[14].strip()
            emp_lname_en = item[15].strip()
            title_en = item[16]
            trm_res_en = item[17]
            emp_birth = item[18].strftime('%d/%m/%Y')
            emp_sex = item[19]
            dept_th = item[20]
            wmonth = item[21]
        
            for col_num in range(len(item)):
                if(col_num==0):
                    ws.write(row_num, 0, counter, font_style)
                elif(col_num==1):
                    ws.write(row_num, col_num, emp_id, font_style)
                elif(col_num==2):
                    ws.write(row_num, col_num, title_th + " " + emp_fname_th + "  " + emp_lname_th, font_style)
                elif(col_num==3):
                    ws.write(row_num, col_num, title_en + " " + emp_fname_en + "  " + emp_lname_en, font_style)
                elif(col_num==4):
                    ws.write(row_num, col_num, emp_sect, font_style)
                elif(col_num==5):
                    ws.write(row_num, col_num, emp_rank, font_style)
                elif(col_num==6):
                    ws.write(row_num, col_num, trm_doc_date, font_style)
                elif(col_num==7):
                    ws.write(row_num, col_num, emp_join_date, font_style)
                elif(col_num==8):
                    ws.write(row_num, col_num, emp_term_date, font_style)
                elif(col_num==9):
                    ws.write(row_num, col_num, wmonth, font_style)
                elif(col_num==10):
                    ws.write(row_num, col_num, sts_th, font_style)
                elif(col_num==11):
                    ws.write(row_num, col_num, trm_res_th, font_style)

            row_num += 1
            counter += 1

        font_style = xlwt.easyxf('font: bold 1, height 180;')
        ws.write(row_num, 0, "TOTAL " + str(counter-1), font_style)

    wb.save(response)
    return response
