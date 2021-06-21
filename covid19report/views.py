from django.shortcuts import render
from django.utils import timezone
from datetime import datetime
from django.conf import settings
from django.http import HttpResponse
from django.http import JsonResponse
from django.contrib.auth.decorators import login_required
from django.contrib.auth.mixins import PermissionRequiredMixin
from django.contrib.auth.decorators import permission_required
from django.utils import translation
from django.utils.translation import ugettext as _
import django.db as db
from django.db import connection
from page.rules import *
from django.utils import timezone
from page.rules import *
from hrms.settings import MEDIA_ROOT
from base64 import b64encode
from docxtpl import DocxTemplate
from docx.shared import Cm, Mm, Pt, Inches
from os import path
from django.http import FileResponse
from docx2pdf import convert
from io import BytesIO
from docx.enum.text import WD_ALIGN_PARAGRAPH


@permission_required('covid19report.can_access_covid_19_report', login_url='/accounts/login/')
def ViewCovid19Report(request):
	user_language = getDefaultLanguage(request.user.username)
	translation.activate(user_language)	

	page_title = settings.PROJECT_NAME
	db_server = settings.DATABASES['default']['HOST']
	project_name = settings.PROJECT_NAME
	project_version = settings.PROJECT_VERSION  
	today_day = timezone.now().strftime('%d')
	today_month = timezone.now().strftime('%m')
	today_year = timezone.now().year
	today_date = str(today_day) + "-" + today_month + "-" + str(today_year)	

	return render(request, 'covid19report/report_by_person.html', {
        'page_title': page_title, 
        'project_name': project_name, 
        'project_version': project_version,
        'db_server': db_server, 
        'today_date': today_date,
        'database': settings.DATABASES['default']['NAME'],
        'host': settings.DATABASES['default']['HOST'],
	})


@permission_required('covid19report.can_access_covid_19_report', login_url='/accounts/login/')
def AjaxCovid19Report(request):    
	today_date = settings.TODAY_DATE.strftime("%d/%m/%Y")
	emp_id = request.POST.get('emp_id')
	get_vaccine_status_option = request.POST.get('get_vaccine_status_option')

	# print("AJAX Debug : ", emp_id)
	# print("AJAX Debug : ", get_vaccine_status_option)
	employee_obj = None
	record = {}	
	message = ""
	full_name = ""
	phone_number = ""


	sql = "select * from covid_employee_vaccine_update where emp_id=" + str(emp_id) + " and get_vaccine_status=" + str(get_vaccine_status_option) + ";"	
	try:                
		cursor = connection.cursor()
		cursor.execute(sql)
		employee_obj = cursor.fetchone()		
	except db.OperationalError as e:
		message = "<b>Error: please send this error to IT team</b><br>" + str(e)
	except db.Error as e:
		message = "<b>Error: please send this error to IT team</b><br>" + str(e)
	finally:
		cursor.close()

	if employee_obj is not None:
		full_name = employee_obj[1]
		phone_number = employee_obj[2]
		
		get_vaccine_status = employee_obj[3]
		get_vaccine_date = employee_obj[4].strftime("%d/%m/%Y")
		get_vaccine_time = employee_obj[4].strftime("%H:00")
		get_vaccine_place = employee_obj[5]
		file_attach_data = b64encode(employee_obj[7]).decode("utf-8")
		file_attach_type = employee_obj[8]

		if get_vaccine_status_option=="1":
			get_vaccine_status_option_text = "นัดหมายเพื่อฉีดวัคซีนข็มที่ 1"
		elif get_vaccine_status_option=="2":
			get_vaccine_status_option_text = "ได้รับการฉีดวัคซีนเข็มที่ 1 เรียบร้อยแล้ว"
		elif get_vaccine_status_option=="3":
			get_vaccine_status_option_text = "นัดหมายเพื่อฉีดวัคซีนข็มที่ 2"
		elif get_vaccine_status_option=="4":
			get_vaccine_status_option_text = "ได้รับการฉีดวัคซีนเข็มที่ 2 เรียบร้อยแล้ว"

		response = JsonResponse(data={        
			"is_error": False,
			"message": "",
			"emp_id": emp_id,
			"full_name": full_name,
			"phone_number": phone_number,
			"get_vaccine_date": get_vaccine_date,
			"get_vaccine_time": get_vaccine_time,
			"get_vaccine_place": get_vaccine_place,
			"get_vaccine_status_option_text": get_vaccine_status_option_text,
			"file_attach_data": file_attach_data,
			"file_attach_type": file_attach_type,
		})
	else:
		response = JsonResponse(data={        
			"is_error": True,
			"message": "ไม่พบข้อมูล",
			"emp_id": emp_id,
			"full_name": full_name,
			"phone_number": phone_number,
		})

	print("full_name : ", full_name)
	print("phone_number : ", phone_number)

	response.status_code = 200
	return response



@permission_required('covid19report.can_access_covid_19_report', login_url='/accounts/login/')
def ViewCovid19ReportByStatus(request):
	user_language = getDefaultLanguage(request.user.username)
	translation.activate(user_language)	

	page_title = settings.PROJECT_NAME
	db_server = settings.DATABASES['default']['HOST']
	project_name = settings.PROJECT_NAME
	project_version = settings.PROJECT_VERSION  
	today_day = timezone.now().strftime('%d')
	today_month = timezone.now().strftime('%m')
	today_year = timezone.now().year
	today_date = str(today_day) + "-" + today_month + "-" + str(today_year)	

	return render(request, 'covid19report/report_by_status.html', {
        'page_title': page_title, 
        'project_name': project_name, 
        'project_version': project_version,
        'db_server': db_server, 
        'today_date': today_date,
        'database': settings.DATABASES['default']['NAME'],
        'host': settings.DATABASES['default']['HOST'],
	})




@permission_required('covid19report.can_access_covid_19_report', login_url='/accounts/login/')
def AjaxReportByStatus(request):    
	today_date = settings.TODAY_DATE.strftime("%d/%m/%Y")
	get_vaccine_status_option = request.POST.get('get_vaccine_status_option')
	
	emp_id_from = request.POST.get('emp_id_from')
	emp_id_to = request.POST.get('emp_id_to')

	emp_type = request.POST.get('emp_type')
	post_id = request.POST.get('post_id')
	
	employee_obj = None
	employee_list = []
	record = {}	
	message = ""

	'''
	if get_vaccine_status_option=="":
		response = JsonResponse(data={        
			"is_error": True,
			"message": "ไม่พบข้อมูล",
			"employee_list": list(employee_list),
		})
		response.status_code = 200
		return response
	'''

	print(emp_id_from, emp_id_to, emp_type, post_id, get_vaccine_status_option)

	sql = "select "
	sql += "emp_id,full_name,phone_number,get_vaccine_status,get_vaccine_date,get_vaccine_place,"
	sql += "file_attach,file_attach_data,file_attach_type,upd_date,upd_by,upd_flag,op1,op2,op3,"
	sql += "op4,op5,opd1,opd2 "
	sql += "from covid_employee_vaccine_update "	
	# sql += "where get_vaccine_status=" + str(get_vaccine_status_option)

	if get_vaccine_status_option=="99":
		if emp_id_from!="":
			sql += " where emp_id>=" + str(emp_id_from)
		else:
			sql += " where emp_id>=0"
		
		if emp_id_to!="":
			sql += " and emp_id<=" + str(emp_id_to)
		else:
			sql += " and emp_id<=999999"

		if emp_type!="":
			sql += " and op3='" + str(emp_type) + "'"

		if post_id!="":
			sql += " and op4='" + str(post_id) + "'"		
	else:
		sql += "where get_vaccine_status=" + str(get_vaccine_status_option)
		if emp_id_from!="":
			sql += " and emp_id>=" + str(emp_id_from)
		else:
			sql += " where emp_id>=0"
		
		if emp_id_to!="":
			sql += " and emp_id<=" + str(emp_id_to)
		else:
			sql += " and emp_id<=999999"

		if emp_type!="":
			sql += " and op3='" + str(emp_type) + "'"

		if post_id!="":
			sql += " and op4='" + str(post_id) + "'"	
	sql += " order by emp_id, get_vaccine_status;"

	print("SQL11 : ", sql)
	#print(emp_id_from, emp_id_to)
	# return False

	try:                
		cursor = connection.cursor()
		cursor.execute(sql)
		employee_obj = cursor.fetchall()		
	except db.OperationalError as e:
		message = "<b>Error: please send this error to IT team</b><br>" + str(e)
	except db.Error as e:
		message = "<b>Error: please send this error to IT team</b><br>" + str(e)
	finally:
		cursor.close()

	if employee_obj is not None:
		for item in employee_obj:
			emp_id = item[0]
			full_name = item[1]
			phone_number = item[2] if item[2] is not None else ""
			get_vaccine_status_option = item[3] if item[3] is not None else 0
			
			# get_vaccine_status_option = item[3]
			# print("get_vaccine_status : ", get_vaccine_status_option)

			get_vaccine_date = item[4].strftime("%d/%m/%Y") if item[4] is not None else ""
			get_vaccine_time = item[4].strftime("%H:00") if item[4] is not None else ""
			get_vaccine_place = item[5] if item[5] is not None else ""
			file_attach = item[6] if item[6] is not None else ""
			
			emp_type = item[14] if item[14] is not None else ""
			post_id = item[15] if item[15] is not None else ""
			post_name = item[13] if item[13] is not None else ""
			zone_name = item[12] if item[12] is not None else ""

			# file_attach_data = b64encode(item[7]).decode("utf-8")
			file_attach_data = b64encode(item[7]).decode("utf-8") if item[7] is not None else ""
			file_attach_type = item[8] if item[8] is not None else ""

			if int(get_vaccine_status_option)==0:
				get_vaccine_status_option_text = "พนักงานที่ยังไม่เคยได้รับการฉีดวัคซีน"
			elif int(get_vaccine_status_option)==1:
				get_vaccine_status_option_text = "นัดหมายเพื่อฉีดวัคซีนข็มที่ 1"
			elif int(get_vaccine_status_option)==2:
				get_vaccine_status_option_text = "ได้รับการฉีดวัคซีนเข็มที่ 1 เรียบร้อยแล้ว"
			elif int(get_vaccine_status_option)==3:
				get_vaccine_status_option_text = "นัดหมายเพื่อฉีดวัคซีนข็มที่ 2"
			elif int(get_vaccine_status_option)==4:
				get_vaccine_status_option_text = "ได้รับการฉีดวัคซีนเข็มที่ 2 เรียบร้อยแล้ว"
			elif int(get_vaccine_status_option)==5:
				get_vaccine_status_option_text = "ยังไม่มีข้อมูลการฉีดวัคซีน"
			else:
				get_vaccine_status_option_text = "Error! : " + str(get_vaccine_status_option)

			# print("DEBUG111 : ", get_vaccine_status_option_text)


			record = {
				"emp_id": emp_id,
				"full_name": full_name,
				"phone_number": phone_number,
				"get_vaccine_status_option": get_vaccine_status_option,
				"get_vaccine_status_option_text": get_vaccine_status_option_text,
				# "get_vaccine_status_option_text": "",
				"get_vaccine_date": get_vaccine_date + " " + get_vaccine_time,
				"get_vaccine_place": get_vaccine_place,
				"file_attach": file_attach,
				"emp_type": emp_type,
				"post_id": post_id,
				"post_name": post_name,
				"zone_name": zone_name,
			}

			employee_list.append(record)

		response = JsonResponse(data={        
			"is_error": False,
			"message": "",
			"employee_list": list(employee_list),
		})
	else:
		response = JsonResponse(data={        
			"is_error": True,
			"message": "ไม่พบข้อมูล",
			"employee_list": list(employee_list),
		})

	response.status_code = 200
	return response


@permission_required('covid19report.can_access_covid_19_report', login_url='/accounts/login/')
def AjaxReportByStatus_bk(request):    
	today_date = settings.TODAY_DATE.strftime("%d/%m/%Y")
	get_vaccine_status_option = request.POST.get('get_vaccine_status_option')


	emp_id = request.POST.get('emp_id')
	emp_type = request.POST.get('emp_type')
	post_id = request.POST.get('post_id')
	
	employee_obj = None
	employee_list = []
	record = {}	
	message = ""


	if get_vaccine_status_option=="":
		response = JsonResponse(data={        
			"is_error": True,
			"message": "ไม่พบข้อมูล",
			"employee_list": list(employee_list),
		})
		response.status_code = 200
		return response


	sql = "select "
	sql += "emp_id,full_name,phone_number,get_vaccine_status,get_vaccine_date,get_vaccine_place,"
	sql += "file_attach,file_attach_data,file_attach_type,upd_date,upd_by,upd_flag,op1,op2,op3,"
	sql += "op4,op5,opd1,opd2 "
	sql += "from covid_employee_vaccine_update "
	sql += "where get_vaccine_status=" + str(get_vaccine_status_option)

	if emp_id!="":
		sql += " and emp_id=" + str(emp_id)
	
	if emp_type!="":
		sql += " and op3='" + str(emp_type) + "'"

	if post_id!="":
		sql += " and op4='" + str(post_id) + "'"
	
	sql += ";"

	# print("SQL11 : ", sql)

	try:                
		cursor = connection.cursor()
		cursor.execute(sql)
		employee_obj = cursor.fetchall()		
	except db.OperationalError as e:
		message = "<b>Error: please send this error to IT team</b><br>" + str(e)
	except db.Error as e:
		message = "<b>Error: please send this error to IT team</b><br>" + str(e)
	finally:
		cursor.close()

	if employee_obj is not None:
		for item in employee_obj:
			emp_id = item[0]
			full_name = item[1]
			phone_number = item[2] if item[2] is not None else ""
			get_vaccine_status = item[3] if item[3] is not None else ""
			get_vaccine_date = item[4].strftime("%d/%m/%Y") if item[4] is not None else ""
			get_vaccine_time = item[4].strftime("%H:00") if item[4] is not None else ""
			get_vaccine_place = item[5] if item[5] is not None else ""
			file_attach = item[6] if item[6] is not None else ""
			
			emp_type = item[14] if item[14] is not None else ""
			post_id = item[15] if item[15] is not None else ""
			post_name = item[13] if item[13] is not None else ""
			zone_name = item[12] if item[12] is not None else ""

			# file_attach_data = b64encode(item[7]).decode("utf-8")
			file_attach_data = b64encode(item[7]).decode("utf-8") if item[7] is not None else ""
			file_attach_type = item[8] if item[8] is not None else ""

			if get_vaccine_status_option=="0":
				get_vaccine_status_option_text = "พนักงานที่ยังไม่เคยได้รับการฉีดวัคซีน"
			if get_vaccine_status_option=="1":
				get_vaccine_status_option_text = "นัดหมายเพื่อฉีดวัคซีนข็มที่ 1"
			elif get_vaccine_status_option=="2":
				get_vaccine_status_option_text = "ได้รับการฉีดวัคซีนเข็มที่ 1 เรียบร้อยแล้ว"
			elif get_vaccine_status_option=="3":
				get_vaccine_status_option_text = "นัดหมายเพื่อฉีดวัคซีนข็มที่ 2"
			elif get_vaccine_status_option=="4":
				get_vaccine_status_option_text = "ได้รับการฉีดวัคซีนเข็มที่ 2 เรียบร้อยแล้ว"

			# print("DEBUG111 : ", get_vaccine_status_option_text)


			record = {
				"emp_id": emp_id,
				"full_name": full_name,
				"phone_number": phone_number,
				"get_vaccine_status_option": get_vaccine_status_option,
				"get_vaccine_status_option_text": get_vaccine_status_option_text,
				"get_vaccine_date": get_vaccine_date + " " + get_vaccine_time,
				"get_vaccine_place": get_vaccine_place,
				"file_attach": file_attach,
				"emp_type": emp_type,
				"post_id": post_id,
				"post_name": post_name,
				"zone_name": zone_name,
			}

			employee_list.append(record)

		response = JsonResponse(data={        
			"is_error": False,
			"message": "",
			"employee_list": list(employee_list),
		})
	else:
		response = JsonResponse(data={        
			"is_error": True,
			"message": "ไม่พบข้อมูล",
			"employee_list": list(employee_list),
		})

	response.status_code = 200
	return response


@permission_required('covid19report.can_access_covid_19_report', login_url='/accounts/login/')
def download_pdf(request, *args, **kwargs):    
	base_url = MEDIA_ROOT + '/covid19/template/'
	template_name = base_url + 'covid19.docx'
	file_name = request.user.username + "_รายงานการฉีดวัคซีนโควิด-19"
	is_error = True	
	error_message = ""
	emp_id = kwargs['emp_id']
	get_vaccine_status_option = kwargs['get_vaccine_status_option']
	employee_info = None

	full_name = ""
	phone_number = ""
	get_vaccine_status_option_text = ""
	get_vaccine_date = ""
	get_vaccine_place = ""

	allowed_file_types = {'JPG','JPEG','PNG','GIF'}

	sql = "select * from covid_employee_vaccine_update where emp_id=" + emp_id + " and get_vaccine_status=" + get_vaccine_status_option + ";"
	print("SQLTEST : ", sql)
	try:        
		cursor = connection.cursor()
		cursor.execute(sql)
		employee_info = cursor.fetchone()
	except db.OperationalError as e:
		is_error = True
		error_message = "<b>Error: please send this error to IT team</b><br>" + str(e)
	except db.Error as e:
		is_error = True
		error_message = "<b>Error: please send this error to IT team</b><br>" + str(e)
	finally:
		cursor.close()

	if employee_info is not None:		
		emp_id = employee_info[0]
		full_name = employee_info[1]
		phone_number = employee_info[2]
		get_vaccine_status_option = employee_info[3]
		
		get_vaccine_date = employee_info[4].strftime("%d-%m-%Y %H:00")
		

		get_vaccine_place = employee_info[5]		
		file_attach_type = employee_info[8]

		if get_vaccine_status_option==1:
			get_vaccine_status_option_text = "นัดหมายเพื่อฉีดวัคซีนข็มที่ 1"
		elif get_vaccine_status_option==2:
			get_vaccine_status_option_text = "ได้รับการฉีดวัคซีนเข็มที่ 1 เรียบร้อยแล้ว"
		elif get_vaccine_status_option==3:
			get_vaccine_status_option_text = "นัดหมายเพื่อฉีดวัคซีนข็มที่ 2"
		elif get_vaccine_status_option==4:
			get_vaccine_status_option_text = "ได้รับการฉีดวัคซีนเข็มที่ 2 เรียบร้อยแล้ว"

		document = DocxTemplate(template_name)
		style = document.styles['Normal']
		font = style.font
		font.name = 'AngsanaUPC'
		font.size = Pt(14)
		
		if file_attach_type!="" :
			if file_attach_type.upper() in allowed_file_types:
				context = {
					"emp_id": emp_id,
					"full_name": full_name,
					"phone_number": phone_number,
					"get_vaccine_status_option_text": get_vaccine_status_option_text,
					"get_vaccine_date": get_vaccine_date,
					"get_vaccine_place": get_vaccine_place,
					"is_file_attached": ""
				}
			else:
				context = {
					"emp_id": emp_id,
					"full_name": full_name,
					"phone_number": phone_number,
					"get_vaccine_status_option_text": get_vaccine_status_option_text,
					"get_vaccine_date": get_vaccine_date,
					"get_vaccine_place": get_vaccine_place,
					"is_file_attached": "ไม่รองรับไฟล์ประเภท " + str(file_attach_type.upper())
				}
		else:
			context = {
				"emp_id": emp_id,
				"full_name": full_name,
				"phone_number": phone_number,
				"get_vaccine_status_option_text": get_vaccine_status_option_text,
				"get_vaccine_date": get_vaccine_date,
				"get_vaccine_place": get_vaccine_place,
				"is_file_attached": "ไม่มีเอกสาร"
			}

		document.render(context)
		

		if file_attach_type!="":
			# allowed_file_types = {'JPG','JPEG','PNG','GIF'}
			if file_attach_type.upper() in allowed_file_types:
				binary_img = BytesIO(employee_info[7])

				# document.add_picture(binary_img, width=Inches(2))
				# last_paragraph = document.paragraphs[-1]
				# last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
				
				tables = document.tables
				p = tables[1].rows[0].cells[0].add_paragraph()
				r = p.add_run()
				r.add_picture(binary_img, width=Inches(3))
				p.alignment = WD_ALIGN_PARAGRAPH.CENTER
				r = p.add_run("\n")

		document.save(MEDIA_ROOT + '/covid19/download/' + file_name + ".docx")    

		docx_file = path.abspath("media\\covid19\\download\\" + file_name + ".docx")
		pdf_file = path.abspath("media\\covid19\\download\\" + file_name + ".pdf")
		convert(docx_file, pdf_file)

		return FileResponse(open(pdf_file, 'rb'), content_type='application/pdf')

